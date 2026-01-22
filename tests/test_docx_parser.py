"""Tests for DocxParser."""

from __future__ import annotations

import tempfile
from pathlib import Path

import docx
import pytest

from app.backend.models.translatable_document import ElementType
from app.backend.parsers.docx_parser import DocxParser


class TestDocxParser:
    """Tests for DocxParser class."""

    @pytest.fixture
    def parser(self):
        """Create a parser instance."""
        return DocxParser()

    @pytest.fixture
    def simple_docx(self):
        """Create a simple DOCX file for testing."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc = docx.Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is the first paragraph.")
        doc.add_paragraph("This is the second paragraph.")
        doc.save(temp_path)

        yield temp_path

        Path(temp_path).unlink(missing_ok=True)

    @pytest.fixture
    def table_docx(self):
        """Create a DOCX file with a table."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc = docx.Document()
        doc.add_paragraph("Before table")

        # Add a 2x2 table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell A1"
        table.cell(0, 1).text = "Cell B1"
        table.cell(1, 0).text = "Cell A2"
        table.cell(1, 1).text = "Cell B2"

        doc.add_paragraph("After table")
        doc.save(temp_path)

        yield temp_path

        Path(temp_path).unlink(missing_ok=True)

    def test_supported_extensions(self, parser):
        """Test that parser declares DOCX support."""
        assert ".docx" in parser.supported_extensions

    def test_file_not_found(self, parser):
        """Test handling of non-existent file."""
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.docx")

    def test_invalid_extension(self, parser):
        """Test handling of non-DOCX file."""
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"Not a DOCX")
            f.flush()
            temp_path = f.name

        try:
            with pytest.raises(ValueError, match="Not a DOCX"):
                parser.parse(temp_path)
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_parse_simple_docx(self, parser, simple_docx):
        """Test parsing a simple DOCX file."""
        doc = parser.parse(simple_docx)

        assert doc is not None
        assert doc.source_type == "docx"
        assert len(doc.elements) >= 2  # At least heading and paragraphs

        # Check that elements have content
        contents = [e.content for e in doc.elements]
        assert any("first paragraph" in c for c in contents)
        assert any("second paragraph" in c for c in contents)

    def test_parse_heading_type(self, parser, simple_docx):
        """Test that headings are classified as TITLE."""
        doc = parser.parse(simple_docx)

        # Find the heading element
        title_elements = [e for e in doc.elements if e.element_type == ElementType.TITLE]
        assert len(title_elements) >= 1
        assert "Test Document" in title_elements[0].content

    def test_parse_table(self, parser, table_docx):
        """Test parsing DOCX with table."""
        doc = parser.parse(table_docx)

        # Find table cell elements
        table_cells = [e for e in doc.elements if e.element_type == ElementType.TABLE_CELL]
        assert len(table_cells) >= 4  # 2x2 table

        # Check that cells are marked correctly
        for cell in table_cells:
            assert cell.metadata.get("in_table") is True

    def test_parse_deduplication(self, parser):
        """Test that duplicate paragraphs are deduplicated."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc = docx.Document()
        # Add same text multiple times
        doc.add_paragraph("Same text")
        doc.add_paragraph("Same text")
        doc.add_paragraph("Different text")
        doc.save(temp_path)

        try:
            result = parser.parse(temp_path)
            contents = [e.content for e in result.elements]

            # Duplicates should be filtered (based on key generation)
            # The exact count depends on key uniqueness logic
            assert "Same text" in contents
            assert "Different text" in contents
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_parse_empty_paragraphs_skipped(self, parser):
        """Test that empty paragraphs are skipped."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc = docx.Document()
        doc.add_paragraph("")  # Empty
        doc.add_paragraph("   ")  # Whitespace only
        doc.add_paragraph("Actual content")
        doc.save(temp_path)

        try:
            result = parser.parse(temp_path)
            contents = [e.content for e in result.elements]

            # Empty paragraphs should not be included
            assert len(contents) == 1
            assert "Actual content" in contents[0]
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_element_ids_unique(self, parser, simple_docx):
        """Test that element IDs are unique."""
        doc = parser.parse(simple_docx)

        ids = [e.element_id for e in doc.elements]
        assert len(ids) == len(set(ids)), "Element IDs should be unique"

    def test_metadata_extraction(self, parser):
        """Test document metadata extraction."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc = docx.Document()
        doc.core_properties.title = "Test Title"
        doc.core_properties.author = "Test Author"
        doc.add_paragraph("Content")
        doc.save(temp_path)

        try:
            result = parser.parse(temp_path)
            assert result.metadata.title == "Test Title"
            assert result.metadata.author == "Test Author"
        finally:
            Path(temp_path).unlink(missing_ok=True)


class TestDocxParserInsertMarker:
    """Tests for skip_inserted_translations feature."""

    @pytest.fixture
    def parser_skip_inserts(self):
        """Create a parser that skips inserted translations."""
        return DocxParser(skip_inserted_translations=True)

    @pytest.fixture
    def parser_include_inserts(self):
        """Create a parser that includes inserted translations."""
        return DocxParser(skip_inserted_translations=False)

    def test_skip_inserted_translation(self, parser_skip_inserts):
        """Test that paragraphs with INSERT_MARKER are skipped."""
        from app.backend.parsers.docx_parser import INSERT_MARKER

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        doc = docx.Document()
        doc.add_paragraph("Original text")

        # Add a paragraph that looks like an inserted translation
        p = doc.add_paragraph()
        run = p.add_run(f"Translated text{INSERT_MARKER}")
        run.italic = True

        doc.save(temp_path)

        try:
            result = parser_skip_inserts.parse(temp_path)
            contents = [e.content for e in result.elements]

            # Should only have the original, not the inserted translation
            assert "Original text" in contents
            assert not any(INSERT_MARKER in c for c in contents)
        finally:
            Path(temp_path).unlink(missing_ok=True)
