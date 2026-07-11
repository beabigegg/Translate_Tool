"""TDD tests for docx-header-footer-collection (BR-115).

Covers:
  - AC-1: native path collects header/footer paragraph text as segments.
  - AC-2: header table (incl. nested table) cells reach the collector via
    the existing `_process_table` walker.
  - AC-3: header/footer paragraph+cell extraction strips `<w:txbxContent>`
    text (Option C) so a header-anchored textbox stays exclusively COM-owned;
    the pre-existing BODY fold-in behavior is unchanged (guards a global
    `_p_text_with_breaks` regression). Also guards the COM call site's
    `include_headers=True` kwarg is unchanged.
  - AC-4: a header/footer part shared across sections via
    `is_linked_to_previous` is collected exactly once (element-identity
    dedup, never `id()`).
  - AC-5: all six per-section slots (header/footer x default/first-page/
    even-page) are traversed.
  - AC-6: header/footer collection runs AFTER the body walk, so body segment
    order/text is identical whether or not the document has header/footer
    content.
  - AC-7: translated header/footer text persists across `doc.save()` +
    reopen (write-back/restore round trip).

All fixtures are built in-test with `python-docx` (mirrors
`tests/test_docx_nested_tables.py`). No test reads `docs/TEST_DOC/`.

Anti-tautology rules (CLAUDE.md):
  - Assert on collected `Segment` content and on the reopened-file / captured
    outgoing payload, never that a helper was called or an attribute was set.
  - AC-4 asserts the shared header text appears EXACTLY ONCE (exact string
    match, not `len(segs)` alone).
  - AC-3's COM call-site guard patches `com_helpers.postprocess_docx_shapes_
    with_word` (the process boundary) and asserts the captured kwarg value.

Collection-time imports: module captured at collection time so
`patch.object` is immune to sys.modules contamination (CLAUDE.md promoted
learnings).
"""

from __future__ import annotations

from unittest.mock import MagicMock, patch

import docx
import pytest
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.shared import Emu

import app.backend.processors.docx_processor as _docx_proc


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_client_mock(**kwargs) -> MagicMock:
    """Return a minimal MagicMock client that satisfies processor health checks."""
    m = MagicMock()
    m.health_check.return_value = (True, "ok")
    m.system_prompt = ""
    m.model_type = "general"
    m._is_translation_dedicated.return_value = False
    m._is_translategemma_model.return_value = False
    for k, v in kwargs.items():
        setattr(m, k, v)
    return m


def _add_textbox_to_paragraph(p) -> None:
    """Append a minimal `<w:txbxContent>` textbox run into paragraph `p`,
    with its own `<w:t>` text "TB_TEXT" — sufficient to probe the
    txbx-strip xpath predicate regardless of the full drawing/shape wrapper
    a real Word textbox uses."""
    r = OxmlElement("w:r")
    drawing = OxmlElement("w:drawing")
    txbx_content = OxmlElement("w:txbxContent")
    tb_p = OxmlElement("w:p")
    tb_r = OxmlElement("w:r")
    tb_t = OxmlElement("w:t")
    tb_t.text = "TB_TEXT"
    tb_r.append(tb_t)
    tb_p.append(tb_r)
    txbx_content.append(tb_p)
    drawing.append(txbx_content)
    r.append(drawing)
    p._p.append(r)


def _fake_translate_texts(texts, targets, src_lang, client, **kwargs):
    """Deterministic SOURCE->TRANSLATED fake for `translate_texts`, keyed
    exactly like the real return: {(tgt, src_text): translated_text}."""
    tmap = {}
    for tgt in targets:
        for txt in texts:
            tmap[(tgt, txt)] = txt + "_TR"
    return tmap, len(texts), 0, False


# ---------------------------------------------------------------------------
# AC-1: native collection — header/footer paragraph text becomes segments
# ---------------------------------------------------------------------------

class TestNativeCollectionLeavesNoSourceText:
    def test_header_and_footer_paragraph_text_collected_as_segments(self):
        doc = docx.Document()
        doc.sections[0].header.is_linked_to_previous = False
        doc.sections[0].header.paragraphs[0].text = "HEADER_SRC_TEXT"
        doc.sections[0].footer.is_linked_to_previous = False
        doc.sections[0].footer.paragraphs[0].text = "FOOTER_SRC_TEXT"

        segs = _docx_proc._collect_docx_segments(doc)
        texts = [s.text for s in segs]
        assert "HEADER_SRC_TEXT" in texts
        assert "FOOTER_SRC_TEXT" in texts

    def test_header_footer_segments_flow_through_translate_docx(self, tmp_path):
        """Integration: header/footer segments actually reach the outgoing
        `translate_texts` payload (not merely collected in isolation)."""
        doc = docx.Document()
        doc.sections[0].header.is_linked_to_previous = False
        doc.sections[0].header.paragraphs[0].text = "HEADER_SRC_TEXT"
        in_path = tmp_path / "in.docx"
        out_path = tmp_path / "out.docx"
        doc.save(str(in_path))

        client = _make_client_mock()
        with patch.object(_docx_proc, "translate_texts") as mock_tt:
            mock_tt.side_effect = _fake_translate_texts
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en", client=client,
                include_headers_shapes_via_com=False,
            )
        sent_texts = mock_tt.call_args[0][0]
        assert "HEADER_SRC_TEXT" in sent_texts


# ---------------------------------------------------------------------------
# AC-2: header table (incl. nested table) collected via existing walker
# ---------------------------------------------------------------------------

class TestHeaderTableAndNestedTableCollected:
    def test_header_table_and_nested_table_cells_collected(self):
        doc = docx.Document()
        doc.sections[0].header.is_linked_to_previous = False
        t = doc.sections[0].header.add_table(rows=1, cols=2, width=Emu(2_000_000))
        t.cell(0, 0).text = "HDR_CELL_A"
        t.cell(0, 1).text = "HDR_CELL_B"
        nested = t.cell(0, 1).add_table(rows=1, cols=1)
        nested.cell(0, 0).text = "HDR_NESTED_CELL"

        segs = _docx_proc._collect_docx_segments(doc)
        cell_texts = {s.text for s in segs if s.kind == "cell"}
        assert cell_texts == {"HDR_CELL_A", "HDR_CELL_B", "HDR_NESTED_CELL"}

        # nested table gets its own table_id group, distinct from the outer table
        by_table = {}
        for s in segs:
            if s.kind == "cell":
                by_table.setdefault(s.table_id, []).append(s.text)
        assert sorted(by_table.values(), key=len) == [["HDR_NESTED_CELL"], ["HDR_CELL_A", "HDR_CELL_B"]]


# ---------------------------------------------------------------------------
# AC-3: txbxContent strip (header/footer only) + COM call-site guard
# ---------------------------------------------------------------------------

class TestTxbxContentStrippedFromHeader:
    def test_header_paragraph_excludes_textbox_text(self):
        doc = docx.Document()
        doc.sections[0].header.is_linked_to_previous = False
        hp = doc.sections[0].header.paragraphs[0]
        hp.text = "HDR_PLAIN"
        _add_textbox_to_paragraph(hp)

        segs = _docx_proc._collect_docx_segments(doc)
        hdr_seg = next(s for s in segs if s.ctx.startswith("HdrFtr") and s.kind == "para")
        assert hdr_seg.text == "HDR_PLAIN"
        assert "TB_TEXT" not in hdr_seg.text

    def test_body_paragraph_textbox_fold_in_unchanged(self):
        """Under amended BR-115 (docx-body-textbox-dedup), the BODY path ALSO
        strips `<w:txbxContent>` text from its `para` segment, exactly like
        the header/footer path above — the textbox text is collected exactly
        once via the dedicated `txbx` segment instead of being folded into
        the paragraph. `_p_text_with_breaks` itself remains unchanged; the
        body walk now threads `_p_text_no_txbx` instead."""
        doc = docx.Document()
        bp = doc.add_paragraph("BODY_PLAIN")
        _add_textbox_to_paragraph(bp)

        segs = _docx_proc._collect_docx_segments(doc)
        body_seg = next(s for s in segs if s.ctx == "Body" and s.kind == "para")
        assert body_seg.text == "BODY_PLAIN"
        assert "TB_TEXT" not in body_seg.text


class TestComCallSiteUnchanged:
    def test_include_headers_kwarg_still_true_at_com_boundary(self, tmp_path):
        doc = docx.Document()
        doc.add_paragraph("Body text")
        in_path = tmp_path / "in.docx"
        out_path = tmp_path / "out.docx"
        doc.save(str(in_path))

        client = _make_client_mock()
        with patch.object(_docx_proc, "translate_texts") as mock_tt, \
             patch.object(_docx_proc, "is_win32com_available", return_value=True), \
             patch.object(_docx_proc, "postprocess_docx_shapes_with_word") as mock_com:
            mock_tt.side_effect = _fake_translate_texts
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en", client=client,
                include_headers_shapes_via_com=True,
            )

        assert mock_com.call_count == 1
        assert mock_com.call_args.kwargs["include_headers"] is True


# ---------------------------------------------------------------------------
# AC-4: linked/shared-part dedup by element identity
# ---------------------------------------------------------------------------

class TestLinkedPartCollectedOnce:
    def test_linked_header_collected_exactly_once(self):
        doc = docx.Document()
        doc.sections[0].header.is_linked_to_previous = False
        doc.sections[0].header.paragraphs[0].text = "SHARED_HEADER_TEXT"
        doc.add_section(WD_SECTION.NEW_PAGE)
        # sanity: the new section's header must actually be linked (proxies
        # the same <w:hdr> element) for this test to exercise the dedup path
        assert doc.sections[1].header.is_linked_to_previous is True

        segs = _docx_proc._collect_docx_segments(doc)
        matches = [s for s in segs if s.text == "SHARED_HEADER_TEXT"]
        assert len(matches) == 1

    def test_linked_header_translated_exactly_once_in_output(self, tmp_path):
        doc = docx.Document()
        doc.sections[0].header.is_linked_to_previous = False
        doc.sections[0].header.paragraphs[0].text = "SHARED_HEADER_TEXT"
        doc.add_section(WD_SECTION.NEW_PAGE)
        in_path = tmp_path / "in.docx"
        out_path = tmp_path / "out.docx"
        doc.save(str(in_path))

        client = _make_client_mock()
        with patch.object(_docx_proc, "translate_texts") as mock_tt:
            mock_tt.side_effect = _fake_translate_texts
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en", client=client,
                include_headers_shapes_via_com=False,
            )
        sent_texts = mock_tt.call_args[0][0]
        assert sent_texts.count("SHARED_HEADER_TEXT") == 1

    def test_linked_header_table_cell_collected_exactly_once(self):
        """The paragraph path has a SEPARATE, pre-existing element-identity
        dedup (`seen_par_keys`, shared across the whole collection call) that
        would mask a broken slot-level `seen_parts` guard for plain paragraph
        text. A header TABLE cell has no such cross-slot guard — `_process_
        table`'s `seen_tc` is a fresh local set per call — so this fixture is
        the one that actually falls through to, and exercises, the slot-level
        `seen_parts` element-identity dedup (AC-4's real correctness guard)."""
        doc = docx.Document()
        doc.sections[0].header.is_linked_to_previous = False
        t = doc.sections[0].header.add_table(rows=1, cols=1, width=Emu(1_000_000))
        t.cell(0, 0).text = "SHARED_HDR_CELL"
        doc.add_section(WD_SECTION.NEW_PAGE)
        assert doc.sections[1].header.is_linked_to_previous is True

        segs = _docx_proc._collect_docx_segments(doc)
        matches = [s for s in segs if s.text == "SHARED_HDR_CELL"]
        assert len(matches) == 1


# ---------------------------------------------------------------------------
# AC-5: all six per-section slots traversed
# ---------------------------------------------------------------------------

class TestAllSixSlotsTraversed:
    def test_each_slot_distinct_marker_text_collected(self):
        doc = docx.Document()
        sec = doc.sections[0]
        slot_names = (
            "header", "footer", "first_page_header",
            "first_page_footer", "even_page_header", "even_page_footer",
        )
        for name in slot_names:
            slot = getattr(sec, name)
            slot.is_linked_to_previous = False
            slot.paragraphs[0].text = f"MARK_{name.upper()}"

        segs = _docx_proc._collect_docx_segments(doc)
        texts = {s.text for s in segs}
        for name in slot_names:
            assert f"MARK_{name.upper()}" in texts, f"slot {name} not collected"


# ---------------------------------------------------------------------------
# AC-6: body order/index stability regardless of header/footer content
# ---------------------------------------------------------------------------

class TestBodyIndicesUnaffectedByHeaderCollection:
    def test_body_segment_order_and_text_identical_with_and_without_header(self):
        doc_no_hdr = docx.Document()
        doc_no_hdr.add_paragraph("Body para 0")
        doc_no_hdr.add_paragraph("Body para 1")
        segs_no_hdr = _docx_proc._collect_docx_segments(doc_no_hdr)
        body_no_hdr = [s.text for s in segs_no_hdr if s.ctx == "Body"]

        doc_hdr = docx.Document()
        doc_hdr.sections[0].header.is_linked_to_previous = False
        doc_hdr.sections[0].header.paragraphs[0].text = "HEADER_MARKER"
        doc_hdr.add_paragraph("Body para 0")
        doc_hdr.add_paragraph("Body para 1")
        segs_hdr = _docx_proc._collect_docx_segments(doc_hdr)
        body_hdr = [s.text for s in segs_hdr if s.ctx == "Body"]

        assert body_no_hdr == ["Body para 0", "Body para 1"]
        assert body_hdr == body_no_hdr

    def test_absolute_segment_zero_is_first_body_paragraph_not_header(self):
        """Real falsifiability guard: filtering by `ctx == "Body"` alone would
        NOT detect header/footer collection running before the body walk (it
        would still return the same body-only sublist regardless of overall
        position). This asserts on the UNFILTERED `segs[0]` — the absolute
        first element in document-collection order, which the `docx:{stem}:
        {idx}` post_translate_hook numbering keys off of (BR-115)."""
        doc = docx.Document()
        doc.sections[0].header.is_linked_to_previous = False
        doc.sections[0].header.paragraphs[0].text = "HEADER_MARKER"
        doc.add_paragraph("Body para 0")
        doc.add_paragraph("Body para 1")

        segs = _docx_proc._collect_docx_segments(doc)
        assert segs[0].text == "Body para 0"
        assert segs[1].text == "Body para 1"

    def test_post_translate_hook_block_ids_index_body_paragraphs_only(self, tmp_path):
        """Integration: `docx:{stem}:{idx}` block ids delivered to
        `post_translate_hook` must number BODY paragraphs 0..N-1; header
        content must not shift or interleave with that numbering."""
        doc = docx.Document()
        doc.sections[0].header.is_linked_to_previous = False
        doc.sections[0].header.paragraphs[0].text = "HEADER_MARKER"
        doc.add_paragraph("Body para 0")
        doc.add_paragraph("Body para 1")
        in_path = tmp_path / "in.docx"
        out_path = tmp_path / "out.docx"
        doc.save(str(in_path))

        captured = []
        client = _make_client_mock()
        with patch.object(_docx_proc, "translate_texts") as mock_tt:
            mock_tt.side_effect = _fake_translate_texts
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en", client=client,
                include_headers_shapes_via_com=False,
                post_translate_hook=lambda tuples: captured.extend(tuples),
            )

        stem = in_path.stem
        by_block_id = {block_id: src for block_id, src, _ in captured}
        assert by_block_id[f"docx:{stem}:0"] == "Body para 0"
        assert by_block_id[f"docx:{stem}:1"] == "Body para 1"


# ---------------------------------------------------------------------------
# AC-7: write-back / restore persists across doc.save() + reopen
# ---------------------------------------------------------------------------

class TestWriteBackPersistsAcrossSave:
    def test_reopened_file_contains_translated_header_and_footer_text(self, tmp_path):
        doc = docx.Document()
        doc.sections[0].header.is_linked_to_previous = False
        doc.sections[0].header.paragraphs[0].text = "HEADER_SOURCE"
        doc.sections[0].footer.is_linked_to_previous = False
        doc.sections[0].footer.paragraphs[0].text = "FOOTER_SOURCE"
        doc.add_paragraph("Body source")
        in_path = tmp_path / "in.docx"
        out_path = tmp_path / "out.docx"
        doc.save(str(in_path))

        client = _make_client_mock()
        with patch.object(_docx_proc, "translate_texts") as mock_tt:
            mock_tt.side_effect = _fake_translate_texts
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en", client=client,
                include_headers_shapes_via_com=False,
            )

        reopened = docx.Document(str(out_path))
        header_texts = [p.text for p in reopened.sections[0].header.paragraphs]
        footer_texts = [p.text for p in reopened.sections[0].footer.paragraphs]
        # inserted translation paragraph carries the text plus a trailing
        # zero-width INSERT_MARKER run (same convention as body inserts).
        assert any(t.startswith("HEADER_SOURCE_TR") for t in header_texts)
        assert any(t.startswith("FOOTER_SOURCE_TR") for t in footer_texts)
