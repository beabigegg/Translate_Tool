"""TDD tests for docx-nested-table-collection (BR-81 amended, BR-113, BR-114).

Covers:
  - AC-1: recursive nested-table collection drops zero characters.
  - AC-2: BR-114 layout-frame reroute (direct paragraphs -> body path;
    nested table still reaches the serializer as its own group).
  - AC-3: BR-81 merged-`<w:tc>` dedup by ELEMENT IDENTITY (never `id()`).
  - AC-4: BR-114 conjunction false-positive guards (never wrongly reroute
    a real data cell).
  - AC-5: MAX_TABLE_NESTING_DEPTH bounded recursion, flatten-and-warn.
  - AC-6: legacy pipe-grid degrade (`JSON_STRUCTURED_TRANSLATION_ENABLED=0`).

All fixtures are built in-test with `python-docx`; no test reads any
external sample-document directory (AC-8).

Anti-tautology rules (CLAUDE.md):
  - Assert on collected `Segment` content and on the outgoing serializer/
    client payload, never on `seen_tc` existing or a flag being set.
  - AC-3 asserts WHICH text sits at WHICH (row, col), not len(segs) alone.
  - AC-6 asserts the empty placeholder in the actual serializer input.

Collection-time imports: modules captured at collection time so
`patch.object` is immune to sys.modules contamination (CLAUDE.md
promoted learnings).
"""

from __future__ import annotations

import json
import logging
from collections import defaultdict
from pathlib import Path
from typing import Dict, List, Tuple
from unittest.mock import MagicMock, patch

import docx
import pytest

from app.backend import config

import app.backend.processors.docx_processor as _docx_proc
import app.backend.utils.table_serializer as _table_ser


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_client_mock(**kwargs) -> MagicMock:
    """Return a minimal MagicMock client that satisfies processor health checks."""
    m = MagicMock()
    m.health_check.return_value = (True, "ok")
    m.system_prompt = ""
    m.model_type = "general"
    m._is_translation_dedicated.return_value = False
    m._is_translategemma_model.return_value = False
    for k, v in kwargs.items():
        setattr(m, k, v)
    return m


def _grid_response(rows: List[List[str]]) -> str:
    """Build a legacy pipe-grid string as a mock LLM response."""
    return "\n".join(" | ".join(row) for row in rows)


def _json_grid_response(mapping: Dict[Tuple[int, int], str]) -> str:
    """Build a coordinate-JSON reply string (BR-79/BR-82) from a
    {(row, col): translation} mapping."""
    return json.dumps({
        "cells": [{"row": r, "col": c, "translation": t} for (r, c), t in mapping.items()]
    })


def _cell_groups_by_table_id(segs) -> Dict[int, list]:
    by_table = defaultdict(list)
    for s in segs:
        if s.kind == "cell":
            by_table[s.table_id].append(s)
    return by_table


# ---------------------------------------------------------------------------
# AC-1: nested-table character parity — nothing silently dropped
# ---------------------------------------------------------------------------

class TestNestedCollectionCharacterParity:
    def test_nested_table_text_not_dropped_recursive_walk(self):
        """Outer 1x1 frame containing an inner 2x2 table (design.md fixture):
        every character across BOTH tables must survive collection, whether
        it lands as a rerouted `para` segment or as a nested `cell` segment.

        Falsifiability: delete the `cell.tables` recursion in `_process_table`
        and the inner table's 4 cell texts vanish from `collected_total`,
        making this assertion fail.
        """
        doc = docx.Document()
        outer = doc.add_table(rows=1, cols=1)
        outer.cell(0, 0).paragraphs[0].text = "Frame prose line one"
        outer.cell(0, 0).add_paragraph("Frame prose line two")
        nested = outer.cell(0, 0).add_table(rows=2, cols=2)
        nested.cell(0, 0).text = "r0c0"
        nested.cell(0, 1).text = "r0c1"
        nested.cell(1, 0).text = "r1c0"
        nested.cell(1, 1).text = "r1c1"

        expected_total = (
            len("Frame prose line one") + len("Frame prose line two")
            + len("r0c0") + len("r0c1") + len("r1c0") + len("r1c1")
        )

        segs = _docx_proc._collect_docx_segments(doc)
        collected_total = sum(len(s.text) for s in segs)
        assert collected_total == expected_total, (
            f"expected {expected_total} chars collected across both tables, "
            f"got {collected_total} (segments: "
            f"{[(s.kind, s.row, s.col, s.table_id, s.text) for s in segs]})"
        )

    def test_seen_tc_survives_id_recycling_60x5(self):
        """Sanity/scale regression for the `seen_tc` amendment (design.md Q4):
        walking a 60x5 table (300 cells, no merges) through `_process_table`
        must retain all 300 distinct texts.

        NOTE on falsifiability: `_process_table`'s dedup happens to be masked
        against the id()-recycling hazard here because every emitted cell is
        immediately retained via `Segment.ref` (the very "undocumented
        invariant" design.md Q3 warns against relying on) — once a cell's
        `Segment` is appended, CPython cannot recycle that cell's address for
        a later, distinct cell, so an `id()`-keyed `seen_tc` would NOT
        reliably go RED on this exact fixture. The genuinely falsifiable
        regression (a code path that does NOT retain cells, so id()
        recycling is observable) is exercised below by
        `TestNestingDepthGuard::test_flatten_at_depth_limit_survives_id_recycling_60x5`,
        which targets `_flatten_nested_table_text` — see that test's
        docstring for the verified-RED evidence.
        """
        doc = docx.Document()
        nrows, ncols = 60, 5
        t = doc.add_table(rows=nrows, cols=ncols)
        expected_texts = set()
        for r in range(nrows):
            for c in range(ncols):
                text = f"r{r}c{c}"
                t.cell(r, c).text = text
                expected_texts.add(text)

        segs = _docx_proc._collect_docx_segments(doc)
        cell_segs = [s for s in segs if s.kind == "cell"]
        assert len(cell_segs) == 300, (
            f"expected 300 distinct cell segments (no merges), got {len(cell_segs)}"
        )
        collected_texts = {s.text for s in cell_segs}
        assert collected_texts == expected_texts


# ---------------------------------------------------------------------------
# AC-2: BR-114 layout-frame reroute
# ---------------------------------------------------------------------------

class TestLayoutFrameReroute:
    def test_frame_cell_direct_paragraphs_routed_to_body_path(self):
        """A cell with a nested table AND full row width (1x1 outer frame)
        reroutes its DIRECT paragraphs to the body/`para` path; the nested
        table still ships as its own `cell`-segment group, and the outer
        cell leaves an empty placeholder (BR-114).
        """
        doc = docx.Document()
        outer = doc.add_table(rows=1, cols=1)
        outer.cell(0, 0).paragraphs[0].text = "Frame prose line one"
        outer.cell(0, 0).add_paragraph("Frame prose line two")
        nested = outer.cell(0, 0).add_table(rows=2, cols=2)
        nested.cell(0, 0).text = "r0c0"
        nested.cell(0, 1).text = "r0c1"
        nested.cell(1, 0).text = "r1c0"
        nested.cell(1, 1).text = "r1c1"

        segs = _docx_proc._collect_docx_segments(doc)

        para_texts = [s.text for s in segs if s.kind == "para"]
        assert "Frame prose line one" in para_texts
        assert "Frame prose line two" in para_texts

        groups = list(_cell_groups_by_table_id(segs).values())
        assert sorted(len(g) for g in groups) == [1, 4], (
            "expected one 1-cell outer group (empty placeholder) and one "
            "4-cell nested group"
        )
        outer_group = next(g for g in groups if len(g) == 1)
        assert outer_group[0].text == "", (
            "rerouted frame cell must leave an EMPTY placeholder cell segment, "
            "not the aggregated frame prose"
        )
        inner_group = next(g for g in groups if len(g) == 4)
        inner_by_pos = {(s.row, s.col): s.text for s in inner_group}
        assert inner_by_pos == {
            (0, 0): "r0c0", (0, 1): "r0c1", (1, 0): "r1c0", (1, 1): "r1c1",
        }

    def test_nested_table_under_frame_cell_reaches_serializer(self, tmp_path):
        """Integration (AC-2): the nested table's group must reach
        `client.translate_json` (the outgoing boundary), not merely be set
        on a `Segment` attribute. The outer 1x1 frame's placeholder cell has
        no content, so it must short-circuit with ZERO whole-table calls
        (only the nested group calls the client)."""
        doc = docx.Document()
        outer = doc.add_table(rows=1, cols=1)
        outer.cell(0, 0).text = "Frame prose"
        nested = outer.cell(0, 0).add_table(rows=2, cols=2)
        nested.cell(0, 0).text = "r0c0"
        nested.cell(0, 1).text = "r0c1"
        nested.cell(1, 0).text = "r1c0"
        nested.cell(1, 1).text = "r1c1"
        in_path = tmp_path / "in.docx"
        doc.save(str(in_path))
        out_path = tmp_path / "out.docx"

        mock_client = _make_client_mock()
        mock_client.translate_json.return_value = (True, _json_grid_response({
            (0, 0): "R0C0_TR", (0, 1): "R0C1_TR", (1, 0): "R1C0_TR", (1, 1): "R1C1_TR",
        }))

        with patch.object(_docx_proc, "translate_texts", return_value=({}, 0, 0, False)):
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en",
                client=mock_client,
                include_headers_shapes_via_com=False,
            )

        assert mock_client.translate_json.call_count == 1, (
            "expected exactly one whole-table JSON call (the nested table's "
            "group); the outer frame's empty placeholder group must short-"
            f"circuit with none — got {mock_client.translate_json.call_count}"
        )
        sent_payload = mock_client.translate_json.call_args[0][0]
        cells_json = sent_payload.split("\n\n", 1)[1]
        sent_cells = {
            (c["row"], c["col"]): c["text"] for c in json.loads(cells_json)["cells"]
        }
        assert sent_cells == {
            (0, 0): "r0c0", (0, 1): "r0c1", (1, 0): "r1c0", (1, 1): "r1c1",
        }


# ---------------------------------------------------------------------------
# AC-3: BR-81 merged-cell dedup by element identity
# ---------------------------------------------------------------------------

class TestMergedCellDedup:
    def test_merged_cell_spanning_columns_emits_single_segment_at_origin(self):
        """A <w:tc> merged across 4 columns must be emitted exactly ONCE, at
        its origin (lowest) column — never once per spanned column."""
        doc = docx.Document()
        t = doc.add_table(rows=2, cols=4)
        t.cell(0, 0).merge(t.cell(0, 3))
        t.cell(0, 0).text = "FRAME"
        for c in range(4):
            t.cell(1, c).text = f"d{c}"

        segs = _docx_proc._collect_docx_segments(doc)
        cell_segs = [s for s in segs if s.kind == "cell"]
        assert len(cell_segs) == 5, (
            f"expected 5 cell segments (1 merged + 4 data), got {len(cell_segs)}"
        )
        by_pos = {(s.row, s.col): s.text for s in cell_segs}
        assert by_pos[(0, 0)] == "FRAME"
        assert (0, 1) not in by_pos and (0, 2) not in by_pos and (0, 3) not in by_pos, (
            "merged cell must NOT be emitted at its spanned columns"
        )
        for c in range(4):
            assert by_pos[(1, c)] == f"d{c}"

    def test_49_distinct_tc_not_52_segments(self):
        """4 rows x 13 cols = 52 naive (row, col) slots; one row-0 merge
        spanning 4 columns collapses 3 duplicate <w:tc> repeats, leaving 49
        DISTINCT elements -> 49 segments (change-request.md 52-vs-49 shape).
        Assert WHICH text sits at WHICH (row, col), not merely the count.
        """
        doc = docx.Document()
        t = doc.add_table(rows=4, cols=13)
        for r in range(4):
            for c in range(13):
                t.cell(r, c).text = f"r{r}c{c}"
        t.cell(0, 0).merge(t.cell(0, 3))
        t.cell(0, 0).text = "MERGED_ROW0"

        segs = _docx_proc._collect_docx_segments(doc)
        cell_segs = [s for s in segs if s.kind == "cell"]
        assert len(cell_segs) == 49, f"expected 49 distinct cell segments, got {len(cell_segs)}"

        by_pos = {(s.row, s.col): s.text for s in cell_segs}
        assert by_pos[(0, 0)] == "MERGED_ROW0"
        for c in (1, 2, 3):
            assert (0, c) not in by_pos, f"merged column {c} must not be re-emitted"
        for r in range(4):
            for c in range(13):
                if r == 0 and c in (0, 1, 2, 3):
                    continue
                assert by_pos[(r, c)] == f"r{r}c{c}"


# ---------------------------------------------------------------------------
# AC-4: BR-114 false-positive guards — the conjunction must never misfire
# ---------------------------------------------------------------------------

class TestFrameRerouteFalsePositiveGuard:
    def test_multi_column_data_table_many_paragraphs_not_rerouted(self):
        """A real multi-column data cell with MANY direct paragraphs but NO
        nested table must stay on the table-cell path — paragraph count is
        explicitly NOT a routing signal (design.md Q1)."""
        doc = docx.Document()
        t = doc.add_table(rows=2, cols=3)
        t.cell(0, 0).paragraphs[0].text = "Para one of data cell"
        t.cell(0, 0).add_paragraph("Para two of data cell")
        t.cell(0, 0).add_paragraph("Para three of data cell")
        t.cell(0, 1).text = "B1"
        t.cell(0, 2).text = "C1"
        t.cell(1, 0).text = "A2"
        t.cell(1, 1).text = "B2"
        t.cell(1, 2).text = "C2"

        segs = _docx_proc._collect_docx_segments(doc)
        para_texts = {s.text for s in segs if s.kind == "para"}
        assert not any("data cell" in txt for txt in para_texts), (
            "a many-paragraph data cell without a nested table must NEVER be "
            "rerouted to the body path"
        )
        cell_by_pos = {(s.row, s.col): s.text for s in segs if s.kind == "cell"}
        assert cell_by_pos[(0, 0)] == (
            "Para one of data cell\nPara two of data cell\nPara three of data cell"
        )

    def test_full_width_cell_without_nested_table_not_rerouted(self):
        """A cell that spans the full row width (a 1x1 table's only cell)
        but has NO nested table must NOT be rerouted — the conjunction's
        second operand alone is not sufficient."""
        doc = docx.Document()
        t = doc.add_table(rows=1, cols=1)
        t.cell(0, 0).text = "Full width, no nested table"

        segs = _docx_proc._collect_docx_segments(doc)
        assert not any(
            s.kind == "para" and s.text == "Full width, no nested table" for s in segs
        )
        cell_segs = [s for s in segs if s.kind == "cell"]
        assert len(cell_segs) == 1
        assert cell_segs[0].row == 0 and cell_segs[0].col == 0
        assert cell_segs[0].text == "Full width, no nested table"

    def test_nested_table_cell_not_full_width_not_rerouted(self):
        """A cell with a nested table but NOT spanning the full row width
        (grid_span=1 in a 3-column table) must NOT be rerouted — the
        conjunction's first operand alone is not sufficient. The nested
        table is still collected as its own group regardless (BR-114)."""
        doc = docx.Document()
        t = doc.add_table(rows=1, cols=3)
        t.cell(0, 0).text = "Cell0 direct text"
        nested = t.cell(0, 0).add_table(rows=1, cols=1)
        nested.cell(0, 0).text = "NESTED_TEXT"
        t.cell(0, 1).text = "Cell1"
        t.cell(0, 2).text = "Cell2"

        segs = _docx_proc._collect_docx_segments(doc)
        assert not any(s.kind == "para" and s.text == "Cell0 direct text" for s in segs), (
            "a nested-but-not-full-width cell must NEVER be rerouted"
        )
        groups = list(_cell_groups_by_table_id(segs).values())
        assert sorted(len(g) for g in groups) == [1, 3], (
            "expected one 3-cell outer group and one 1-cell nested group"
        )
        outer_group = next(g for g in groups if len(g) == 3)
        outer_by_pos = {(s.row, s.col): s.text for s in outer_group}
        assert outer_by_pos == {
            (0, 0): "Cell0 direct text", (0, 1): "Cell1", (0, 2): "Cell2",
        }
        nested_group = next(g for g in groups if len(g) == 1)
        assert nested_group[0].text == "NESTED_TEXT"


# ---------------------------------------------------------------------------
# AC-5: MAX_TABLE_NESTING_DEPTH bounded recursion, flatten-and-warn
# ---------------------------------------------------------------------------

class TestNestingDepthGuard:
    def test_recursion_terminates_at_max_depth_with_single_warning(self, caplog):
        """Nest MAX_TABLE_NESTING_DEPTH + 2 (5) levels deep in single-cell
        tables. Depths 1-3 must each form their own group; depth-3's cell
        must fold depths 4-5's text into its own cell text (never drop),
        logging EXACTLY ONE warning, and recursion must terminate (test
        completing at all proves no hang)."""
        assert config.MAX_TABLE_NESTING_DEPTH == 3, "test assumes the documented default"

        doc = docx.Document()
        t1 = doc.add_table(rows=1, cols=1)
        t2 = t1.cell(0, 0).add_table(rows=1, cols=1)
        t3 = t2.cell(0, 0).add_table(rows=1, cols=1)
        t4 = t3.cell(0, 0).add_table(rows=1, cols=1)
        t5 = t4.cell(0, 0).add_table(rows=1, cols=1)
        t5.cell(0, 0).text = "DEEPEST_TEXT"

        with caplog.at_level(logging.WARNING, logger="TranslateTool"):
            segs = _docx_proc._collect_docx_segments(doc)

        warning_records = [
            r for r in caplog.records
            if r.name == "TranslateTool" and "MAX_TABLE_NESTING_DEPTH" in r.getMessage()
        ]
        assert len(warning_records) == 1, (
            f"expected exactly one depth-limit WARNING, got {len(warning_records)}"
        )

        cell_segs = [s for s in segs if s.kind == "cell"]
        table_ids = {s.table_id for s in cell_segs}
        assert len(table_ids) == 3, (
            f"expected 3 independent table groups (depths 1-3), got {len(table_ids)}"
        )

        deepest_created_tid = max(table_ids)
        depth3_segs = [s for s in cell_segs if s.table_id == deepest_created_tid]
        assert len(depth3_segs) == 1
        assert "DEEPEST_TEXT" in depth3_segs[0].text, (
            "depth-4/5 text must be flattened into the depth-3 cell's text, "
            "never dropped"
        )

    def test_flatten_at_depth_limit_survives_id_recycling_60x5(self):
        """Falsifiability guard for the `_flatten_nested_table_text` dedup
        amendment (design.md Q4/BR-81): the flatten-at-depth-limit path is
        the genuinely id()-recycling-prone spot in this module, because it
        does NOT retain any `_Cell` object after extracting its text (unlike
        `_process_table`, whose emitted `Segment.ref` keeps a cell alive for
        the rest of the walk, which happens to mask the hazard there). This
        test nests a wide 60x5 table one level past MAX_TABLE_NESTING_DEPTH
        so it is folded via `_flatten_nested_table_text`, then asserts every
        one of its 300 distinct cell texts survived into the flattened text.

        Verified RED: temporarily keying `_flatten_nested_table_text`'s
        `seen_tc_local` by `id(cell._tc)` instead of the element collapses
        this to only 8/300 surviving texts (confirmed by direct execution
        against a scratch copy of docx_processor.py, matching
        evidence/probe_id_recycling.py's measured 8-distinct-of-300 result;
        the swap was reverted before this file was finalized).
        """
        assert config.MAX_TABLE_NESTING_DEPTH == 3, "test assumes the documented default"

        doc = docx.Document()
        t1 = doc.add_table(rows=1, cols=1)
        t2 = t1.cell(0, 0).add_table(rows=1, cols=1)
        t3 = t2.cell(0, 0).add_table(rows=1, cols=1)
        nrows, ncols = 60, 5
        t4 = t3.cell(0, 0).add_table(rows=nrows, cols=ncols)  # depth 4: over the limit
        expected_tokens = set()
        for r in range(nrows):
            for c in range(ncols):
                text = f"r{r}c{c}"
                t4.cell(r, c).text = text
                expected_tokens.add(text)

        segs = _docx_proc._collect_docx_segments(doc)
        cell_segs = [s for s in segs if s.kind == "cell"]
        table_ids = {s.table_id for s in cell_segs}
        assert len(table_ids) == 3, "depth-4 must be flattened, not its own group"

        depth3_segs = [s for s in cell_segs if s.table_id == max(table_ids)]
        assert len(depth3_segs) == 1
        flattened_tokens = set(depth3_segs[0].text.split("\n"))
        missing = expected_tokens - flattened_tokens
        assert not missing, (
            f"{len(missing)}/300 nested cell texts were dropped during depth-limit "
            f"flattening (e.g. {sorted(missing)[:5]}) — consistent with the "
            "id()-recycling false-collision hazard if seen_tc_local regressed "
            "to id(cell._tc)"
        )


# ---------------------------------------------------------------------------
# AC-6: legacy pipe-grid degrade (JSON_STRUCTURED_TRANSLATION_ENABLED=0)
# ---------------------------------------------------------------------------

class TestLegacyPipeGridDegrade:
    def test_legacy_flag_off_nested_table_own_grid_no_crash(self, tmp_path, monkeypatch):
        """Flag OFF: a nested table degrades to its OWN small dense grid via
        the unchanged serialize()/parse(), never merged into the parent's
        matrix — both tables translate independently with no crash."""
        monkeypatch.setattr(config, "JSON_STRUCTURED_TRANSLATION_ENABLED", False)

        doc = docx.Document()
        outer = doc.add_table(rows=2, cols=2)
        nested = outer.cell(0, 0).add_table(rows=2, cols=2)
        nested.cell(0, 0).text = "N00"
        nested.cell(0, 1).text = "N01"
        nested.cell(1, 0).text = "N10"
        nested.cell(1, 1).text = "N11"
        outer.cell(0, 1).text = "B1"
        outer.cell(1, 0).text = "B2"
        outer.cell(1, 1).text = "B3"
        in_path = tmp_path / "in.docx"
        doc.save(str(in_path))
        out_path = tmp_path / "out.docx"

        mock_client = _make_client_mock()
        mock_client.translate_once.side_effect = [
            (True, _grid_response([["", "OUT_B1_TR"], ["OUT_B2_TR", "OUT_B3_TR"]])),
            (True, _grid_response([["N00_TR", "N01_TR"], ["N10_TR", "N11_TR"]])),
        ]

        with patch.object(_docx_proc, "translate_texts", return_value=({}, 0, 0, False)):
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en",
                client=mock_client,
                include_headers_shapes_via_com=False,
                output_mode="replace",
            )

        assert mock_client.translate_once.call_count == 2, (
            "expected one legacy whole-table call per table group (outer + "
            f"nested), got {mock_client.translate_once.call_count}"
        )

        result_doc = docx.Document(str(out_path))
        result_outer = result_doc.tables[0]
        assert result_outer.cell(0, 1).text == "OUT_B1_TR"
        assert result_outer.cell(1, 0).text == "OUT_B2_TR"
        assert result_outer.cell(1, 1).text == "OUT_B3_TR"

        result_nested = result_outer.cell(0, 0).tables[0]
        assert result_nested.cell(0, 0).text == "N00_TR"
        assert result_nested.cell(0, 1).text == "N01_TR"
        assert result_nested.cell(1, 0).text == "N10_TR"
        assert result_nested.cell(1, 1).text == "N11_TR"

    def test_legacy_flag_off_rerouted_frame_leaves_placeholder(self, tmp_path, monkeypatch):
        """Flag OFF + BR-114 reroute: the outer 1x1 frame's direct paragraph
        goes to the paragraph path; its cell-segment placeholder is EMPTY
        and must reach the legacy serializer's `serialize()` input as such
        (not an intermediate var) — while the nested table still translates
        independently."""
        monkeypatch.setattr(config, "JSON_STRUCTURED_TRANSLATION_ENABLED", False)

        doc = docx.Document()
        outer = doc.add_table(rows=1, cols=1)
        outer.cell(0, 0).text = "Frame prose"
        nested = outer.cell(0, 0).add_table(rows=2, cols=2)
        nested.cell(0, 0).text = "N00"
        nested.cell(0, 1).text = "N01"
        nested.cell(1, 0).text = "N10"
        nested.cell(1, 1).text = "N11"
        in_path = tmp_path / "in.docx"
        doc.save(str(in_path))
        out_path = tmp_path / "out.docx"

        mock_client = _make_client_mock()
        mock_client.translate_once.side_effect = [
            (True, ""),  # outer 1x1 all-empty placeholder group: unparseable, falls to no-op
            (True, _grid_response([["N00_TR", "N01_TR"], ["N10_TR", "N11_TR"]])),
        ]

        real_serialize = _table_ser.serialize
        with patch.object(_docx_proc, "translate_texts", return_value=({}, 0, 0, False)), \
             patch("app.backend.utils.table_serializer.serialize", wraps=real_serialize) as spy_serialize:
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en",
                client=mock_client,
                include_headers_shapes_via_com=False,
                output_mode="replace",
            )

        assert spy_serialize.call_count == 2
        outer_sent_cells = list(spy_serialize.call_args_list[0][0][0])
        assert len(outer_sent_cells) == 1
        assert outer_sent_cells[0].row == 0
        assert outer_sent_cells[0].col == 0
        assert outer_sent_cells[0].content == "", (
            "the rerouted frame cell's placeholder must reach serialize() as "
            "an EMPTY content cell, not the frame prose"
        )

        result_doc = docx.Document(str(out_path))
        result_outer = result_doc.tables[0]
        result_nested = result_outer.cell(0, 0).tables[0]
        assert result_nested.cell(0, 0).text == "N00_TR"
        assert result_nested.cell(0, 1).text == "N01_TR"
        assert result_nested.cell(1, 0).text == "N10_TR"
        assert result_nested.cell(1, 1).text == "N11_TR"
