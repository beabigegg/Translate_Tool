"""TDD tests for docx-body-textbox-dedup (BR-115 amended).

Covers:
  - AC-1/AC-7: a body paragraph hosting a `<w:txbxContent>` textbox collects
    the textbox text EXACTLY ONCE (via `_txbx_iter_texts`), never folded into
    the paragraph's own `para` segment. This is the RED->GREEN boundary test.
  - AC-2: same isolation for a table-cell paragraph hosting a textbox.
  - AC-3: full collect->translate->restore round trip; the textbox
    translation lands only inside `<w:txbxContent>`, never in the enclosing
    paragraph/cell body runs.
  - AC-4: extractor-family consistency — `_p_text_no_txbx` is the extractor
    actually invoked both at collection time (body walk) and at the three
    restore-time re-read sites (SDT branch, cell branch, tail-scan), proven
    by spying on the module-level callable, not by inspecting source text.
  - AC-5: textbox-free body/cell collection is unaffected.
  - AC-6: `_txbx_iter_texts` itself still extracts a multi-paragraph
    textbox's own content in full (unaffected by this change).

All fixtures are built in-test with `python-docx`. No `docs/TEST_DOC/`
fixtures are read. `_add_textbox_to_paragraph` is reused (not duplicated)
from `tests/test_docx_header_footer.py` per the implementation plan.

Anti-tautology rules (CLAUDE.md):
  - AC-1/AC-2 assert SELECTION (which segment holds the textbox text), not
    just segment counts.
  - AC-3 asserts the translated text lands in the correct XML location
    (inside vs. outside `<w:txbxContent>`), not merely that translation
    happened somewhere.
  - AC-4 spies on the actual `_p_text_no_txbx` invocation across a genuine
    resume-idempotency round trip, framed honestly as extractor-family
    hygiene (the three restore-site swaps are inert to output).

Collection-time imports: modules captured at collection time so
`patch.object` is immune to sys.modules contamination (CLAUDE.md promoted
learnings).
"""

from __future__ import annotations

from unittest.mock import MagicMock, patch

import docx
import pytest
from docx.oxml import OxmlElement

import app.backend.processors.docx_processor as _docx_proc
from app.backend.processors.docx_processor import Segment
from tests.test_docx_header_footer import _add_textbox_to_paragraph


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_client_mock(**kwargs) -> MagicMock:
    """Return a minimal MagicMock client that satisfies processor health checks."""
    m = MagicMock()
    m.health_check.return_value = (True, "ok")
    m.system_prompt = ""
    m.model_type = "general"
    m._is_translation_dedicated.return_value = False
    m._is_translategemma_model.return_value = False
    for k, v in kwargs.items():
        setattr(m, k, v)
    return m


def _fake_translate_texts(texts, targets, src_lang, client, **kwargs):
    """Deterministic SOURCE->TRANSLATED fake for `translate_texts`, keyed
    exactly like the real return: {(tgt, src_text): translated_text}."""
    tmap = {}
    for tgt in targets:
        for txt in texts:
            tmap[(tgt, txt)] = txt + "_TR"
    return tmap, len(texts), 0, False


def _add_multi_paragraph_textbox(host_paragraph, lines) -> None:
    """Like `_add_textbox_to_paragraph` but with N paragraphs of textbox
    content, to probe `_txbx_iter_texts`'s own multi-paragraph join (AC-6)."""
    r = OxmlElement("w:r")
    drawing = OxmlElement("w:drawing")
    txbx_content = OxmlElement("w:txbxContent")
    for line in lines:
        tb_p = OxmlElement("w:p")
        tb_r = OxmlElement("w:r")
        tb_t = OxmlElement("w:t")
        tb_t.text = line
        tb_r.append(tb_t)
        tb_p.append(tb_r)
        txbx_content.append(tb_p)
    drawing.append(txbx_content)
    r.append(drawing)
    host_paragraph._p.append(r)


def _txt_nodes_outside_txbx(root):
    return root.xpath(
        ".//*[local-name()='t'][not(ancestor::*[local-name()='txbxContent'])]"
    )


def _txt_nodes_inside_txbx(root):
    return root.xpath(".//*[local-name()='txbxContent']//*[local-name()='t']")


# ---------------------------------------------------------------------------
# AC-1 / AC-2 / AC-7: body/cell paragraph excludes textbox text at collection
# ---------------------------------------------------------------------------

class TestBodyTextboxCollectedOnce:
    def test_body_paragraph_excludes_textbox_text(self):
        """RED->GREEN boundary (AC-7). Pre-fix: docx_processor.py L427 threads
        the default `_p_text_with_breaks` into the body walk, so the textbox
        text folds into the paragraph's own `para` segment
        ("BODY_PLAINTB_TEXT"). Post-fix (`text_extractor=_p_text_no_txbx`):
        the `para` segment excludes it entirely and the textbox's own text
        is collected exactly once via the `txbx` segment.
        """
        doc = docx.Document()
        bp = doc.add_paragraph("BODY_PLAIN")
        _add_textbox_to_paragraph(bp)

        segs = _docx_proc._collect_docx_segments(doc)

        body_seg = next(s for s in segs if s.ctx == "Body" and s.kind == "para")
        assert body_seg.text == "BODY_PLAIN"
        assert "TB_TEXT" not in body_seg.text

        txbx_segs = [s for s in segs if s.kind == "txbx"]
        assert len(txbx_segs) == 1
        assert txbx_segs[0].text == "TB_TEXT"

    def test_cell_paragraph_excludes_textbox_text(self):
        """AC-2: same isolation for a table-cell paragraph hosting a textbox."""
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "CELL_PLAIN"
        _add_textbox_to_paragraph(cell.paragraphs[0])

        segs = _docx_proc._collect_docx_segments(doc)

        cell_seg = next(s for s in segs if s.kind == "cell")
        assert cell_seg.text == "CELL_PLAIN"
        assert "TB_TEXT" not in cell_seg.text

        txbx_segs = [s for s in segs if s.kind == "txbx"]
        assert len(txbx_segs) == 1
        assert txbx_segs[0].text == "TB_TEXT"


# ---------------------------------------------------------------------------
# AC-3: restore round trip isolates the textbox translation
# ---------------------------------------------------------------------------

class TestRestoreIsolatesTextboxTranslation:
    def test_body_textbox_translation_isolated_from_paragraph_body(self, tmp_path):
        doc = docx.Document()
        bp = doc.add_paragraph("BODY_PLAIN")
        _add_textbox_to_paragraph(bp)
        in_path = tmp_path / "in.docx"
        out_path = tmp_path / "out.docx"
        doc.save(str(in_path))

        client = _make_client_mock()
        with patch.object(_docx_proc, "translate_texts") as mock_tt:
            mock_tt.side_effect = _fake_translate_texts
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en", client=client,
                include_headers_shapes_via_com=False,
            )

        sent_texts = mock_tt.call_args[0][0]
        assert "BODY_PLAIN" in sent_texts
        assert "TB_TEXT" in sent_texts
        assert "BODY_PLAINTB_TEXT" not in sent_texts

        reopened = docx.Document(str(out_path))
        body_el = reopened.element.body
        inside = "".join(t.text or "" for t in _txt_nodes_inside_txbx(body_el))
        outside = "".join(t.text or "" for t in _txt_nodes_outside_txbx(body_el))
        assert "TB_TEXT_TR" in inside
        assert "TB_TEXT_TR" not in outside
        assert "BODY_PLAIN_TR" in outside

    def test_cell_textbox_translation_isolated_from_cell_body(self, tmp_path):
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "CELL_PLAIN"
        _add_textbox_to_paragraph(cell.paragraphs[0])
        in_path = tmp_path / "in.docx"
        out_path = tmp_path / "out.docx"
        doc.save(str(in_path))

        client = _make_client_mock()
        client.translate_json.return_value = (
            True,
            '{"cells": [{"row": 0, "col": 0, "translation": "CELL_PLAIN_TR"}]}',
        )
        with patch.object(_docx_proc, "translate_texts") as mock_tt:
            mock_tt.side_effect = _fake_translate_texts
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en", client=client,
                include_headers_shapes_via_com=False,
            )

        sent_texts = mock_tt.call_args[0][0]
        assert "TB_TEXT" in sent_texts
        assert "CELL_PLAINTB_TEXT" not in sent_texts

        sent_payload = client.translate_json.call_args[0][0]
        assert "CELL_PLAIN" in sent_payload
        assert "TB_TEXT" not in sent_payload

        reopened = docx.Document(str(out_path))
        body_el = reopened.element.body
        inside = "".join(t.text or "" for t in _txt_nodes_inside_txbx(body_el))
        outside = "".join(t.text or "" for t in _txt_nodes_outside_txbx(body_el))
        assert "TB_TEXT_TR" in inside
        assert "TB_TEXT_TR" not in outside
        assert "CELL_PLAIN_TR" in outside


# ---------------------------------------------------------------------------
# AC-4: extractor-family consistency (collection + restore-time re-reads)
# ---------------------------------------------------------------------------

class TestExtractorFamilyConsistency:
    def test_body_collection_invokes_stripping_extractor(self):
        """Collection-side identity: the body walk must actually CALL
        `_p_text_no_txbx`, not merely accept it as an unused parameter."""
        doc = docx.Document()
        bp = doc.add_paragraph("BODY_PLAIN")
        _add_textbox_to_paragraph(bp)

        with patch.object(
            _docx_proc, "_p_text_no_txbx", wraps=_docx_proc._p_text_no_txbx
        ) as spy:
            _docx_proc._collect_docx_segments(doc)

        assert spy.call_count >= 1, (
            "body collection did not invoke _p_text_no_txbx at all"
        )

    def test_sdt_restore_reread_invokes_stripping_extractor(self):
        """Restore-hygiene (SDT branch, ~L550): the existing-translation
        re-read on a resume pass must use `_p_text_no_txbx`."""
        doc = docx.Document()
        sdt = OxmlElement("w:sdt")
        sdt_content = OxmlElement("w:sdtContent")
        sdt_p = OxmlElement("w:p")
        sdt_r = OxmlElement("w:r")
        sdt_t = OxmlElement("w:t")
        sdt_t.text = "SDT_SRC"
        sdt_r.append(sdt_t)
        sdt_p.append(sdt_r)
        sdt_content.append(sdt_p)
        sdt.append(sdt_content)
        doc._body._body.append(sdt)

        seg = Segment("para", sdt, "Body > SDT-Placeholder", "SDT_SRC")
        tmap = {("vi", "SDT_SRC", None): "SDT_SRC_TR"}

        # First pass inserts the translation (append mode).
        _docx_proc._insert_docx_translations(
            doc, [seg], tmap, targets=["vi"], output_mode="append"
        )

        # Second pass (resume idempotency) re-reads the existing translation.
        with patch.object(
            _docx_proc, "_p_text_no_txbx", wraps=_docx_proc._p_text_no_txbx
        ) as spy:
            _docx_proc._insert_docx_translations(
                doc, [seg], tmap, targets=["vi"], output_mode="append"
            )

        assert spy.call_count >= 1, (
            "SDT-branch restore re-read did not invoke _p_text_no_txbx"
        )

    def test_cell_restore_reread_invokes_stripping_extractor(self):
        """Restore-hygiene (table-cell branch, ~L596): a `para`-kind segment
        whose paragraph parent is a `_Cell` (the BR-114 reroute case) takes
        the cell-restore re-read path, distinct from the plain-paragraph
        tail-scan path below."""
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "CELL_SRC"
        p = cell.paragraphs[0]
        assert isinstance(p._parent, _docx_proc._Cell)

        seg = Segment("para", p, "Body > Tbl(r0,c0)", "CELL_SRC")
        tmap = {("vi", "CELL_SRC", None): "CELL_SRC_TR"}

        _docx_proc._insert_docx_translations(
            doc, [seg], tmap, targets=["vi"], output_mode="append"
        )

        with patch.object(
            _docx_proc, "_p_text_no_txbx", wraps=_docx_proc._p_text_no_txbx
        ) as spy:
            _docx_proc._insert_docx_translations(
                doc, [seg], tmap, targets=["vi"], output_mode="append"
            )

        assert spy.call_count >= 1, (
            "cell-branch restore re-read did not invoke _p_text_no_txbx"
        )

    def test_tail_scan_restore_reread_invokes_stripping_extractor(self):
        """Restore-hygiene (plain-paragraph tail-scan, `_scan_our_tail_texts`
        called at ~L664)."""
        doc = docx.Document()
        p = doc.add_paragraph("PARA_SRC")

        seg = Segment("para", p, "Body", "PARA_SRC")
        tmap = {("vi", "PARA_SRC", None): "PARA_SRC_TR"}

        _docx_proc._insert_docx_translations(
            doc, [seg], tmap, targets=["vi"], output_mode="append"
        )

        with patch.object(
            _docx_proc, "_p_text_no_txbx", wraps=_docx_proc._p_text_no_txbx
        ) as spy:
            _docx_proc._insert_docx_translations(
                doc, [seg], tmap, targets=["vi"], output_mode="append"
            )

        assert spy.call_count >= 1, (
            "plain-paragraph tail-scan restore re-read did not invoke _p_text_no_txbx"
        )


# ---------------------------------------------------------------------------
# AC-5: textbox-free body/cell collection unaffected
# ---------------------------------------------------------------------------

class TestTextboxFreeBodyUnchanged:
    def test_textbox_free_body_and_cell_segments_unchanged(self):
        doc = docx.Document()
        doc.add_paragraph("Plain body paragraph, no textbox.")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Plain cell, no textbox."

        segs = _docx_proc._collect_docx_segments(doc)

        body_seg = next(s for s in segs if s.ctx == "Body" and s.kind == "para")
        assert body_seg.text == "Plain body paragraph, no textbox."

        cell_seg = next(s for s in segs if s.kind == "cell")
        assert cell_seg.text == "Plain cell, no textbox."

        assert not any(s.kind == "txbx" for s in segs)


# ---------------------------------------------------------------------------
# AC-6: `_txbx_iter_texts` own multi-paragraph extraction unaffected
# ---------------------------------------------------------------------------

class TestTextboxOwnContentUnaffected:
    def test_txbx_iter_texts_extracts_full_multiparagraph_textbox_content(self):
        doc = docx.Document()
        bp = doc.add_paragraph("HOST")
        _add_multi_paragraph_textbox(bp, ["LINE_ONE", "LINE_TWO", "LINE_THREE"])

        results = list(_docx_proc._txbx_iter_texts(doc))

        assert len(results) == 1
        _tx, text = results[0]
        assert text == "LINE_ONE\nLINE_TWO\nLINE_THREE"
