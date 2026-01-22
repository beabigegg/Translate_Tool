"""Tests for InlineRenderer."""

from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import MagicMock

import pytest

from app.backend.models.translatable_document import (
    DocumentMetadata,
    ElementType,
    PageInfo,
    TranslatableDocument,
    TranslatableElement,
)
from app.backend.renderers.base import RenderMode
from app.backend.renderers.inline_renderer import InlineRenderer


class TestInlineRenderer:
    """Tests for InlineRenderer class."""

    @pytest.fixture
    def renderer(self):
        """Create a renderer instance."""
        return InlineRenderer(font_size_pt=10, italic=True)

    @pytest.fixture
    def sample_document(self):
        """Create a sample TranslatableDocument."""
        return TranslatableDocument(
            source_path="/test/document.pdf",
            source_type="pdf",
            elements=[
                TranslatableElement(
                    element_id="elem_1",
                    content="Hello, world!",
                    element_type=ElementType.TEXT,
                    page_num=1,
                    should_translate=True,
                ),
                TranslatableElement(
                    element_id="elem_2",
                    content="This is a test.",
                    element_type=ElementType.TEXT,
                    page_num=1,
                    should_translate=True,
                ),
                TranslatableElement(
                    element_id="elem_3",
                    content="Page two content.",
                    element_type=ElementType.TEXT,
                    page_num=2,
                    should_translate=True,
                ),
            ],
            pages=[
                PageInfo(page_num=1, width=612, height=792),
                PageInfo(page_num=2, width=612, height=792),
            ],
            metadata=DocumentMetadata(page_count=2),
        )

    def test_supported_modes(self, renderer):
        """Test that renderer only supports INLINE mode."""
        assert RenderMode.INLINE in renderer.supported_modes
        assert len(renderer.supported_modes) == 1

    def test_output_extension(self, renderer):
        """Test output file extension."""
        assert renderer.output_extension == ".docx"

    def test_render_invalid_mode(self, renderer, sample_document):
        """Test that render fails with non-INLINE mode."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        try:
            with pytest.raises(ValueError, match="only supports INLINE"):
                renderer.render(
                    sample_document,
                    temp_path,
                    translations={"Hello, world!": "你好，世界！"},
                    mode=RenderMode.OVERLAY,
                )
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_render_basic(self, renderer, sample_document):
        """Test basic rendering to DOCX."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        try:
            translations = {
                "Hello, world!": "你好，世界！",
                "This is a test.": "這是一個測試。",
                "Page two content.": "第二頁內容。",
            }

            renderer.render(
                sample_document,
                temp_path,
                translations=translations,
                mode=RenderMode.INLINE,
            )

            # Verify file was created
            assert Path(temp_path).exists()
            assert Path(temp_path).stat().st_size > 0

            # Read back and verify content
            import docx

            doc = docx.Document(temp_path)

            # Should have page headers and content
            paragraphs = [p.text for p in doc.paragraphs]
            assert "-- Page 1 --" in paragraphs
            assert "Hello, world!" in paragraphs
            assert "你好，世界！" in [p.text.replace("\u200b", "") for p in doc.paragraphs]
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_render_missing_translation(self, renderer, sample_document):
        """Test rendering with missing translations."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        try:
            # Only provide partial translations
            translations = {
                "Hello, world!": "你好，世界！",
            }

            renderer.render(
                sample_document,
                temp_path,
                translations=translations,
                mode=RenderMode.INLINE,
            )

            # Verify file was created
            assert Path(temp_path).exists()

            # Read back and check for missing translation placeholder
            import docx

            doc = docx.Document(temp_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            assert "[Translation missing]" in all_text
        finally:
            Path(temp_path).unlink(missing_ok=True)

    def test_render_non_translatable_elements(self, renderer):
        """Test rendering with non-translatable elements."""
        document = TranslatableDocument(
            source_path="/test/document.pdf",
            source_type="pdf",
            elements=[
                TranslatableElement(
                    element_id="elem_1",
                    content="Header text",
                    element_type=ElementType.HEADER,
                    page_num=1,
                    should_translate=False,  # Not translatable
                ),
                TranslatableElement(
                    element_id="elem_2",
                    content="Body text",
                    element_type=ElementType.TEXT,
                    page_num=1,
                    should_translate=True,
                ),
            ],
            pages=[PageInfo(page_num=1, width=612, height=792)],
            metadata=DocumentMetadata(page_count=1),
        )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        try:
            renderer.render(
                document,
                temp_path,
                translations={"Body text": "正文文字"},
                mode=RenderMode.INLINE,
            )

            assert Path(temp_path).exists()
        finally:
            Path(temp_path).unlink(missing_ok=True)


class TestInlineRendererFromSegments:
    """Tests for render_from_segments backward compatibility method."""

    @pytest.fixture
    def renderer(self):
        """Create a renderer instance."""
        return InlineRenderer(font_size_pt=10, italic=True)

    def test_render_from_segments_paragraph(self, renderer):
        """Test inserting translation after paragraph."""
        import docx

        doc = docx.Document()
        p = doc.add_paragraph("Original text")

        segments = [
            {"kind": "para", "ref": p, "text": "Original text"},
        ]
        translations = {("zh-TW", "Original text"): "原文翻譯"}

        ok_cnt, skip_cnt = renderer.render_from_segments(
            doc, segments, translations, targets=["zh-TW"]
        )

        assert ok_cnt == 1
        assert skip_cnt == 0

    def test_render_from_segments_no_translation(self, renderer):
        """Test skipping segments without translation."""
        import docx

        doc = docx.Document()
        p = doc.add_paragraph("Original text")

        segments = [
            {"kind": "para", "ref": p, "text": "Original text"},
        ]
        translations = {}  # No translations

        ok_cnt, skip_cnt = renderer.render_from_segments(
            doc, segments, translations, targets=["zh-TW"]
        )

        assert ok_cnt == 0
        assert skip_cnt == 1
