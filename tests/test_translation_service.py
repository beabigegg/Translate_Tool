"""Non-regression tests for translation_service.py — AC-6 (table-context-translation).

Tests that the non-table paragraph translation path is unaffected by the
table-context-translation changes. Calls translate_docx() directly (NOT via
translate_document() — avoiding the wrong-entry-point tautology).

Anti-tautology rule: these tests assert the RESULT of paragraph translation
(actual text in the output document), not just call counts or wiring.

Collection-time imports (CLAUDE.md patch.object rule):
  All LLM boundary mocks use patch.object with collection-time references.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List, Optional
from unittest.mock import MagicMock, patch

import pytest

import app.backend.processors.docx_processor as _docx_proc
import app.backend.services.translation_service as _ts


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx_paragraphs_only(tmp_path: Path, paragraphs: List[str]) -> Path:
    """Create a minimal DOCX with only body paragraphs (no tables)."""
    import docx as _docx_lib
    doc = _docx_lib.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    path = tmp_path / "para_only.docx"
    doc.save(str(path))
    return path


def _mock_client() -> MagicMock:
    m = MagicMock()
    m.health_check.return_value = (True, "ok")
    m.system_prompt = ""
    m.model_type = "general"
    m._is_translation_dedicated.return_value = False
    m._is_translategemma_model.return_value = False
    return m


# ---------------------------------------------------------------------------
# AC-6: Non-table paragraph path unaffected
# ---------------------------------------------------------------------------

class TestNonTableParagraphUnaffected:
    """Verify that non-table paragraph translation continues to work correctly (AC-6)."""

    def test_non_table_paragraph_translation_unaffected(self, tmp_path):
        """A DOCX with only body paragraphs should still be translated correctly.

        After IP-4: paragraphs use (tgt, text, None) tmap key internally;
        the OUTPUT behavior (translation appended) must be unchanged vs baseline.
        """
        src_text = "Hello world"
        expected_translation = "你好世界"

        in_path = _make_docx_paragraphs_only(tmp_path, [src_text])
        out_path = tmp_path / "out.docx"

        # translate_texts returns the expected translation for the paragraph
        tmap = {("zh", src_text): expected_translation}

        with patch.object(_docx_proc, "translate_texts", return_value=(tmap, 1, 0, False)):
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en",
                client=_mock_client(),
                include_headers_shapes_via_com=False,
            )

        import docx as _docx_lib
        out_doc = _docx_lib.Document(str(out_path))
        all_texts = [p.text for p in out_doc.paragraphs]

        # Source text must still be present (append mode)
        assert any(src_text in t for t in all_texts), (
            f"Source paragraph missing from output: {all_texts}"
        )
        # Translation must be appended
        assert any(expected_translation in t for t in all_texts), (
            f"Expected translation '{expected_translation}' not found in output: {all_texts}"
        )

    def test_paragraph_tmap_key_col_none(self, tmp_path):
        """After IP-4: paragraphs are re-keyed to (tgt, text, None) before restore pass.

        translate_texts() returns standard 2-element keys; translate_docx() re-keys
        them to (tgt, text, None) so _insert_docx_translations can use a uniform
        3-element key for both para (col=None) and table (col=<int>) segments.

        This test verifies the full round-trip: 2-element key from translate_texts
        → re-keyed to (tgt, text, None) → restore pass applies the translation.
        """
        src_text = "Test paragraph"
        expected = "测试段落"

        in_path = _make_docx_paragraphs_only(tmp_path, [src_text])
        out_path = tmp_path / "out.docx"

        def fake_translate_texts(texts, targets, src_lang, client, **kwargs):
            # Standard 2-element keys as returned by the real translate_texts()
            result = {}
            for tgt in targets:
                for text in texts:
                    result[(tgt, text)] = expected
            return result, len(texts), 0, False

        with patch.object(_docx_proc, "translate_texts", side_effect=fake_translate_texts):
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en",
                client=_mock_client(),
                include_headers_shapes_via_com=False,
            )

        import docx as _docx_lib
        out_doc = _docx_lib.Document(str(out_path))
        all_texts = [p.text for p in out_doc.paragraphs]

        assert any(expected in t for t in all_texts), (
            f"Expected translation '{expected}' not in output (col=None tmap key not matched): "
            f"{all_texts}"
        )

    def test_multiple_paragraphs_all_translated_via_non_table_path(self, tmp_path):
        """Multiple paragraphs are all translated correctly via the paragraph path."""
        paras = ["First sentence", "Second sentence", "Third sentence"]
        translations = {
            "First sentence": "第一句",
            "Second sentence": "第二句",
            "Third sentence": "第三句",
        }

        in_path = _make_docx_paragraphs_only(tmp_path, paras)
        out_path = tmp_path / "out.docx"

        def fake_translate(texts, targets, src_lang, client, **kwargs):
            # translate_texts returns 2-element keys (tgt, text) — NOT 3-element.
            # translate_docx re-keys them to (tgt, text, None) internally.
            result = {}
            for tgt in targets:
                for text in texts:
                    if text in translations:
                        result[(tgt, text)] = translations[text]
            return result, len(texts), 0, False

        with patch.object(_docx_proc, "translate_texts", side_effect=fake_translate):
            _docx_proc.translate_docx(
                str(in_path), str(out_path),
                targets=["zh"], src_lang="en",
                client=_mock_client(),
                include_headers_shapes_via_com=False,
            )

        import docx as _docx_lib
        out_doc = _docx_lib.Document(str(out_path))
        all_texts = [p.text for p in out_doc.paragraphs]

        for t in translations.values():
            assert any(t in txt for txt in all_texts), (
                f"Translation '{t}' missing from output: {all_texts}"
            )
