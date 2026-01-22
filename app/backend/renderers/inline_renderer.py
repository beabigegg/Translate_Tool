"""Inline renderer for inserting translations after original text.

This renderer produces DOCX output where translations are inserted
as new paragraphs following the original text blocks.
"""

from __future__ import annotations

import logging
from typing import TYPE_CHECKING, Callable, Dict, List, Optional

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from app.backend.config import INSERT_FONT_SIZE_PT
from app.backend.renderers.base import BaseRenderer, RenderMode

if TYPE_CHECKING:
    from app.backend.models.translatable_document import TranslatableDocument

logger = logging.getLogger(__name__)

# Marker used to identify inserted translation paragraphs
INSERT_MARKER = "\u200b"


class InlineRenderer(BaseRenderer):
    """Renderer that inserts translations inline after original text.

    This is the default rendering mode for DOCX output. Each translation
    is inserted as a new paragraph after the original text, styled in
    italic with a configurable font size.
    """

    def __init__(
        self,
        font_size_pt: int = INSERT_FONT_SIZE_PT,
        italic: bool = True,
        log: Callable[[str], None] = lambda s: None,
    ):
        """Initialize the inline renderer.

        Args:
            font_size_pt: Font size for inserted translations.
            italic: Whether to italicize translations.
            log: Logging callback function.
        """
        self.font_size_pt = font_size_pt
        self.italic = italic
        self.log = log

    @property
    def supported_modes(self) -> list[RenderMode]:
        """Supported rendering modes."""
        return [RenderMode.INLINE]

    @property
    def output_extension(self) -> str:
        """Output file extension."""
        return ".docx"

    def render(
        self,
        document: "TranslatableDocument",
        output_path: str,
        translations: Dict[str, str],
        mode: RenderMode = RenderMode.INLINE,
    ) -> None:
        """Render translated document to DOCX with inline translations.

        Args:
            document: Source TranslatableDocument.
            output_path: Path for output DOCX file.
            translations: Mapping from original text to translated text.
            mode: Must be RenderMode.INLINE.

        Raises:
            ValueError: If mode is not INLINE.
        """
        if mode != RenderMode.INLINE:
            raise ValueError(
                f"InlineRenderer only supports INLINE mode, got {mode.value}"
            )

        # Create new DOCX document
        doc = docx.Document()

        # Track current page for page headers
        current_page = 0

        # Get elements in reading order
        elements = document.get_elements_in_reading_order()

        for element in elements:
            # Add page separator when page changes
            if element.page_num != current_page:
                current_page = element.page_num
                doc.add_heading(f"-- Page {current_page} --", level=1)

            # Skip non-translatable elements
            if not element.should_translate:
                # Still add non-translatable content but without translation
                if element.content.strip():
                    p = doc.add_paragraph(element.content.strip())
                    # Mark as non-translated
                    if p.runs:
                        p.runs[0].font.color.rgb = docx.shared.RGBColor(128, 128, 128)
                continue

            original_text = element.content.strip()
            if not original_text:
                continue

            # Add original text
            doc.add_paragraph(original_text)

            # Add translation if available
            if original_text in translations:
                translated = translations[original_text]
                self._add_translation_paragraph(doc, translated)
            else:
                # Add placeholder for missing translation
                self._add_translation_paragraph(
                    doc, f"[Translation missing] {original_text[:50]}..."
                )

        # Save document
        doc.save(output_path)
        self.log(f"[Renderer] Saved inline output: {output_path}")

    def _add_translation_paragraph(self, doc: docx.Document, text: str) -> Paragraph:
        """Add a translation paragraph with standard formatting.

        Args:
            doc: Target DOCX document.
            text: Translation text to add.

        Returns:
            The created Paragraph object.
        """
        p = doc.add_paragraph()
        lines = text.split("\n")

        for i, line in enumerate(lines):
            run = p.add_run(line)
            if self.italic:
                run.italic = True
            if self.font_size_pt:
                run.font.size = Pt(self.font_size_pt)
            if i < len(lines) - 1:
                run.add_break()

        # Add marker to identify our inserted paragraphs
        marker_run = p.add_run(INSERT_MARKER)
        if self.italic:
            marker_run.italic = True
        if self.font_size_pt:
            marker_run.font.size = Pt(self.font_size_pt)

        return p

    def render_from_segments(
        self,
        doc: docx.Document,
        segments: List[dict],
        translations: Dict[str, str],
        targets: List[str],
    ) -> tuple[int, int]:
        """Render translations into an existing DOCX document.

        This method is for backward compatibility with the existing
        DOCX processing flow.

        Args:
            doc: Existing DOCX document to modify.
            segments: List of segment dictionaries with 'kind', 'ref', 'text' keys.
            translations: Mapping from (target, text) to translation.
            targets: List of target languages.

        Returns:
            Tuple of (inserted_count, skipped_count).
        """
        ok_cnt = 0
        skip_cnt = 0

        for seg in segments:
            text = seg.get("text", "")
            if not text.strip():
                continue

            # Check if any translation exists
            has_translation = any((tgt, text) in translations for tgt in targets)
            if not has_translation:
                self.log(f"[SKIP] No translation: {text[:50]}...")
                skip_cnt += 1
                continue

            # Get all translations for this segment
            trans_list = []
            for tgt in targets:
                if (tgt, text) in translations:
                    trans_list.append(translations[(tgt, text)])
                else:
                    trans_list.append(f"[Translation missing|{tgt}] {text[:50]}...")

            # Insert translations based on segment kind
            kind = seg.get("kind", "para")
            ref = seg.get("ref")

            if kind == "para" and ref is not None:
                self._insert_after_paragraph(ref, trans_list)
                ok_cnt += 1
            elif kind == "txbx" and ref is not None:
                self._insert_into_textbox(ref, trans_list)
                ok_cnt += 1
            else:
                self.log(f"[WARN] Unknown segment kind: {kind}")
                skip_cnt += 1

        self.log(f"[Renderer] Inserted: {ok_cnt}, Skipped: {skip_cnt}")
        return ok_cnt, skip_cnt

    def _insert_after_paragraph(
        self,
        p: Paragraph,
        translations: List[str],
    ) -> None:
        """Insert translation paragraphs after a source paragraph.

        Args:
            p: Source paragraph to insert after.
            translations: List of translations to insert.
        """
        anchor = p
        for trans in translations:
            new_p_element = OxmlElement("w:p")
            anchor._p.addnext(new_p_element)
            new_p = Paragraph(new_p_element, p._parent)

            lines = trans.split("\n")
            for i, line in enumerate(lines):
                run = new_p.add_run(line)
                if self.italic:
                    run.italic = True
                if self.font_size_pt:
                    run.font.size = Pt(self.font_size_pt)
                if i < len(lines) - 1:
                    run.add_break()

            # Add marker
            marker_run = new_p.add_run(INSERT_MARKER)
            if self.italic:
                marker_run.italic = True
            if self.font_size_pt:
                marker_run.font.size = Pt(self.font_size_pt)

            anchor = new_p

    def _insert_into_textbox(
        self,
        txbx_element,
        translations: List[str],
    ) -> None:
        """Insert translation paragraphs into a text box.

        Args:
            txbx_element: Text box XML element.
            translations: List of translations to insert.
        """
        for trans in translations:
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")

            if self.italic:
                rpr.append(OxmlElement("w:i"))
            if self.font_size_pt:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), str(int(self.font_size_pt * 2)))
                rpr.append(sz)

            r.append(rpr)

            lines = trans.split("\n")
            for i, line in enumerate(lines):
                if i > 0:
                    r.append(OxmlElement("w:br"))
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = line
                r.append(t)

            # Add marker
            marker_t = OxmlElement("w:t")
            marker_t.set(qn("xml:space"), "preserve")
            marker_t.text = INSERT_MARKER
            r.append(marker_t)

            p.append(r)
            txbx_element.append(p)
