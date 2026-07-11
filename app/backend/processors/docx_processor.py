"""DOCX translation processor."""

from __future__ import annotations

import os
import threading
from pathlib import Path
from typing import Any, Callable, Dict, Iterator, List, Optional, Tuple, TYPE_CHECKING

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.backend import config
from app.backend.clients.ollama_client import OllamaClient
from app.backend.config import (
    DEFAULT_MAX_BATCH_CHARS,
    INSERT_FONT_SIZE_PT,
    MAX_SEGMENTS,
    MAX_TEXT_LENGTH,
)
from app.backend.processors.com_helpers import is_win32com_available, postprocess_docx_shapes_with_word
from app.backend.services.translation_service import translate_texts
from app.backend.utils import json_translation
from app.backend.utils.exceptions import ApiError, check_document_size_limits
from app.backend.utils.length_guard import is_suspiciously_short
from app.backend.utils.logging_utils import logger
from app.backend.utils.text_utils import has_cjk, is_numeric_cell, normalize_text, should_translate

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

INSERT_MARKER = "\u200b"


def _p_text_with_breaks(p: Paragraph) -> str:
    parts = []
    for node in p._p.xpath(".//*[local-name()='t' or local-name()='br' or local-name()='tab']"):
        tag = node.tag.split("}", 1)[-1]
        if tag == "t":
            parts.append(node.text or "")
        elif tag == "br":
            parts.append("\n")
        else:
            parts.append(" ")
    return "".join(parts).strip()


def _p_text_no_txbx(p: Paragraph) -> str:
    """Same as `_p_text_with_breaks` but excludes text under `<w:txbxContent>`
    (BR-115, amended). This is the uniform extractor for the whole native DOCX
    collection surface — header/footer, body, and table-cell paragraph/cell
    extraction — so a paragraph or cell hosting a textbox stays exclusively
    owned by its dedicated `_txbx_iter_texts` collection path instead of being
    folded into the enclosing paragraph/cell segment and translated twice.
    Header-anchored textboxes additionally stay exclusively owned by the
    Windows COM shapes pass (`postprocess_docx_shapes_with_word`) on that OS.
    """
    parts = []
    for node in p._p.xpath(
        ".//*[(local-name()='t' or local-name()='br' or local-name()='tab') "
        "and not(ancestor::*[local-name()='txbxContent'])]"
    ):
        tag = node.tag.split("}", 1)[-1]
        if tag == "t":
            parts.append(node.text or "")
        elif tag == "br":
            parts.append("\n")
        else:
            parts.append(" ")
    return "".join(parts).strip()


def _append_after(p: Paragraph, text_block: str, italic: bool = True, font_size_pt: int = INSERT_FONT_SIZE_PT) -> Paragraph:
    new_p = OxmlElement("w:p")
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    lines = text_block.split("\n")
    for i, line in enumerate(lines):
        run = np.add_run(line)
        if italic:
            run.italic = True
        if font_size_pt:
            run.font.size = Pt(font_size_pt)
        if i < len(lines) - 1:
            run.add_break()
    tag = np.add_run(INSERT_MARKER)
    if italic:
        tag.italic = True
    if font_size_pt:
        tag.font.size = Pt(font_size_pt)
    return np


def _is_our_insert_block(p: Paragraph) -> bool:
    return any(INSERT_MARKER in (r.text or "") for r in p.runs)


def _find_last_inserted_after(p: Paragraph, limit: int = 8) -> Optional[Paragraph]:
    ptr = p._p.getnext()
    last = None
    steps = 0
    while ptr is not None and steps < limit:
        if ptr.tag.endswith("}p"):
            q = Paragraph(ptr, p._parent)
            if _is_our_insert_block(q):
                last = q
                steps += 1
                ptr = ptr.getnext()
                continue
        break
    return last


def _scan_our_tail_texts(p: Paragraph, limit: int = 8) -> List[str]:
    ptr = p._p.getnext()
    out = []
    steps = 0
    while ptr is not None and steps < limit:
        if ptr.tag.endswith("}p"):
            q = Paragraph(ptr, p._parent)
            if _is_our_insert_block(q):
                out.append(_p_text_no_txbx(q))
                steps += 1
                ptr = ptr.getnext()
                continue
        break
    return out


def _txbx_iter_texts(doc: Any) -> Iterator[Tuple[Any, str]]:
    def _p_text_flags(p_el):
        parts = []
        for node in p_el.xpath(".//*[local-name()='t' or local-name()='br' or local-name()='tab']"):
            tag = node.tag.split("}", 1)[-1]
            if tag == "t":
                parts.append(node.text or "")
            elif tag == "br":
                parts.append("\n")
            else:
                parts.append(" ")
        text = "".join(parts)
        has_zero = INSERT_MARKER in text
        runs = p_el.xpath(".//*[local-name()='r']")
        visible, ital = [], []
        for run in runs:
            rt = "".join([(t.text or "") for t in run.xpath(".//*[local-name()='t']")])
            if (rt or "").strip():
                visible.append(rt)
                ital.append(bool(run.xpath(".//*[local-name()='i']")))
        all_italic = len(visible) > 0 and all(ital)
        return text, has_zero, all_italic

    for tx in doc._element.xpath(".//*[local-name()='txbxContent']"):
        kept = []
        for p in tx.xpath(".//*[local-name()='p']"):
            text, has_zero, _ = _p_text_flags(p)
            if not (text or "").strip():
                continue
            if has_zero:
                continue
            for line in text.split("\n"):
                if line.strip():
                    kept.append(line.strip())
        if kept:
            joined = "\n".join(kept)
            yield tx, joined


def _txbx_append_paragraph(tx: Any, text_block: str, italic: bool = True, font_size_pt: int = INSERT_FONT_SIZE_PT) -> None:
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if italic:
        rpr.append(OxmlElement("w:i"))
    if font_size_pt:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size_pt * 2)))
        rpr.append(sz)
    r.append(rpr)
    lines = text_block.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            r.append(OxmlElement("w:br"))
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = line
        r.append(t)
    tag = OxmlElement("w:t")
    tag.set(qn("xml:space"), "preserve")
    tag.text = INSERT_MARKER
    r.append(tag)
    p.append(r)
    tx.append(p)


def _txbx_tail_equals(tx: Any, translations: List[str]) -> bool:
    paras = tx.xpath("./*[local-name()='p']")
    if len(paras) < len(translations):
        return False
    tail = paras[-len(translations):]
    for q, expect in zip(tail, translations):
        parts = []
        for node in q.xpath(".//*[local-name()='t' or local-name()='br']"):
            tag = node.tag.split("}", 1)[-1]
            parts.append("\n" if tag == "br" else (node.text or ""))
        if normalize_text("".join(parts).strip()) != normalize_text(expect):
            return False
    return True


class Segment:
    """Represents a segment of text to be translated in a document.

    For table cells (kind="cell"):
        col       — 0-based column index (used as tmap dedup key; BR-81)
        row       — 0-based row index (used for serializer grid building)
        table_id  — monotonically increasing per-document counter assigned in
                    document order (top-level and nested tables alike), used
                    for per-table grouping (BR-113). NOT id() of the table
                    XML element — see BR-113/BR-81 for the lxml proxy-address
                    recycling hazard that makes any id()-derived key unsafe.

    For non-table segments (kind="para", "txbx"):
        col, row, table_id are all None (tmap key uses col=None; BR-81).
    """

    def __init__(
        self,
        kind: str,
        ref: Any,
        ctx: str,
        text: str,
        col: Optional[int] = None,
        row: Optional[int] = None,
        table_id: Optional[int] = None,
    ) -> None:
        self.kind = kind
        self.ref = ref
        self.ctx = ctx
        self.text = text
        self.col = col
        self.row = row
        self.table_id = table_id


def _collect_docx_segments(
    doc: Any,
    max_segments: int = MAX_SEGMENTS,
    max_text_length: int = MAX_TEXT_LENGTH,
) -> List[Segment]:
    segs: List[Segment] = []
    seen_par_keys = set()
    total_text_length = 0
    next_table_id = 0
    depth_limit_warned = False

    def _add_paragraph(p: Paragraph, ctx: str, text_extractor: Callable[[Paragraph], str] = _p_text_with_breaks) -> None:
        nonlocal total_text_length
        try:
            p_key = p._p  # element identity dedup (BR-81 amended: never id())
            if p_key in seen_par_keys:
                return
            txt = text_extractor(p)
            if txt.strip() and not _is_our_insert_block(p):
                segs.append(Segment("para", p, ctx, txt))
                seen_par_keys.add(p_key)
                total_text_length += len(txt)
        except Exception as exc:
            logger.warning("Paragraph processing error: %s", exc)

    def _cell_direct_text(cell: _Cell, text_extractor: Callable[[Paragraph], str] = _p_text_with_breaks) -> str:
        """Aggregate a cell's DIRECT paragraph text only (not nested tables)."""
        parts = []
        for p in cell.paragraphs:
            try:
                txt = text_extractor(p)
                if txt.strip() and not _is_our_insert_block(p):
                    parts.append(txt)
            except Exception:
                pass
        return "\n".join(parts)

    def _flatten_nested_table_text(table: Table, text_extractor: Callable[[Paragraph], str] = _p_text_with_breaks) -> str:
        """Aggregate ALL text within `table`, recursing into further nested
        tables without bound. Used only at the MAX_TABLE_NESTING_DEPTH limit
        (BR-113 flatten-and-warn) to fold over-deep content into the parent
        cell's ordinary cell text instead of dropping it or emitting a group.
        Merged `<w:tc>` repeats are deduped by element identity (BR-81).
        """
        seen_tc_local = set()
        parts = []
        for row in table.rows:
            for cell in row.cells:
                if cell._tc in seen_tc_local:
                    continue
                seen_tc_local.add(cell._tc)
                direct_text = _cell_direct_text(cell, text_extractor)
                if direct_text:
                    parts.append(direct_text)
                for nested in cell.tables:
                    nested_text = _flatten_nested_table_text(nested, text_extractor)
                    if nested_text:
                        parts.append(nested_text)
        return "\n".join(parts)

    def _process_table(table: Table, ctx: str, depth: int, text_extractor: Callable[[Paragraph], str] = _p_text_with_breaks) -> None:
        """Collect one `<w:tbl>` as its own group (own counter-assigned
        table_id, own private 0-based (row, col) space; BR-113). Handles the
        BR-81 merged-cell dedup, the BR-114 frame-reroute gate, and bounded
        recursion into cell.tables with flatten-and-warn at the depth limit.
        """
        nonlocal total_text_length, next_table_id, depth_limit_warned
        next_table_id += 1
        tid = next_table_id
        seen_tc: set = set()
        col_count = table._tbl.col_count
        for r_idx, row in enumerate(table.rows):  # 0-based
            for c_idx, cell in enumerate(row.cells):  # 0-based
                if cell._tc in seen_tc:
                    # Horizontally-merged <w:tc>: row.cells repeats the SAME
                    # element once per spanned column — emit once, at origin
                    # (lowest) column only (BR-81 clarified).
                    continue
                seen_tc.add(cell._tc)

                cell_ctx = f"{ctx} > Tbl(r{r_idx},c{c_idx})"
                has_nested = len(cell.tables) > 0
                recurse_nested = has_nested and depth < config.MAX_TABLE_NESTING_DEPTH
                flatten_needed = has_nested and not recurse_nested

                flatten_extra = ""
                if flatten_needed:
                    flat_parts = []
                    for nested in cell.tables:
                        nested_text = _flatten_nested_table_text(nested, text_extractor)
                        if nested_text:
                            flat_parts.append(nested_text)
                    flatten_extra = "\n".join(flat_parts)
                    if not depth_limit_warned:
                        logger.warning(
                            "DOCX nested table exceeds MAX_TABLE_NESTING_DEPTH=%d "
                            "at %s; flattening its text into the enclosing cell "
                            "instead of collecting it as its own table group",
                            config.MAX_TABLE_NESTING_DEPTH, cell_ctx,
                        )
                        depth_limit_warned = True

                # BR-114: reroute a layout-frame cell's DIRECT paragraphs to the
                # body path iff it has a nested table AND spans the full row
                # width. Never reroute otherwise (paragraph count is not a
                # signal).
                reroute = has_nested and cell.grid_span == col_count

                if reroute:
                    for p in cell.paragraphs:
                        _add_paragraph(p, cell_ctx, text_extractor)
                    # Empty-text placeholder keeps the legacy pipe-grid's
                    # positional (row, col) slot filled (BR-114 flag-off
                    # clause); fold in any depth-limit flatten text so nothing
                    # is silently dropped.
                    placeholder_text = flatten_extra
                    segs.append(Segment(
                        "cell", cell, cell_ctx, placeholder_text,
                        col=c_idx, row=r_idx, table_id=tid,
                    ))
                    total_text_length += len(placeholder_text)
                else:
                    cell_text = _cell_direct_text(cell, text_extractor)
                    if flatten_extra:
                        cell_text = "\n".join(
                            part for part in (cell_text, flatten_extra) if part
                        )
                    # Collect cell even if empty (positional placeholder for serializer)
                    segs.append(Segment(
                        "cell", cell, cell_ctx, cell_text,
                        col=c_idx, row=r_idx, table_id=tid,
                    ))
                    total_text_length += len(cell_text)

                if recurse_nested:
                    # Reading order within a cell: direct-paragraph content
                    # first (already emitted above), then nested tables in
                    # document order (design.md "Recursion / reading order").
                    for nested in cell.tables:
                        _process_table(nested, f"{cell_ctx} > Nested", depth + 1, text_extractor)

    def _process_container_content(container, ctx: str, depth: int = 1, text_extractor: Callable[[Paragraph], str] = _p_text_with_breaks) -> None:
        if container._element is None:
            return
        for child_element in container._element:
            qname = child_element.tag
            if qname.endswith("}p"):
                p = Paragraph(child_element, container)
                _add_paragraph(p, ctx, text_extractor)
            elif qname.endswith("}tbl"):
                table = Table(child_element, container)
                _process_table(table, ctx, depth, text_extractor)
            elif qname.endswith("}sdt"):
                sdt_ctx = f"{ctx} > SDT"
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                placeholder_texts = []
                for t in child_element.xpath(".//w:placeholder//w:t", namespaces=ns):
                    if t.text:
                        placeholder_texts.append(t.text)
                if placeholder_texts:
                    full_placeholder = "".join(placeholder_texts).strip()
                    if full_placeholder:
                        segs.append(Segment("para", child_element, f"{sdt_ctx}-Placeholder", full_placeholder))
                list_items = []
                for item in child_element.xpath(".//w:dropDownList/w:listItem", namespaces=ns):
                    display_text = item.get(qn("w:displayText"))
                    if display_text:
                        list_items.append(display_text)
                if list_items:
                    items_as_text = "\n".join(list_items)
                    segs.append(Segment("para", child_element, f"{sdt_ctx}-Dropdown", items_as_text))
                sdt_content_element = child_element.find(qn("w:sdtContent"))
                if sdt_content_element is not None:
                    class SdtContentWrapper:
                        def __init__(self, element, parent):
                            self._element = element
                            self._parent = parent
                    sdt_content_wrapper = SdtContentWrapper(sdt_content_element, container)
                    _process_container_content(sdt_content_wrapper, sdt_ctx, depth, text_extractor)

    _process_container_content(doc._body, "Body", 1, text_extractor=_p_text_no_txbx)

    for tx, s in _txbx_iter_texts(doc):
        if s.strip() and (has_cjk(s) or should_translate(s, "auto")):
            segs.append(Segment("txbx", tx, "TextBox", s))
            total_text_length += len(s)

    # BR-115: native header/footer collection, appended AFTER the body walk so
    # body segment indices 0..N-1 and docx:{stem}:{idx} hook numbering stay
    # unaffected (AC-6). Dedup by <w:hdr>/<w:ftr> ELEMENT identity (never
    # id() — BR-81/BR-113/BR-115); a linked slot shares the same element with
    # a prior section, so holding the element in a set collects/writes it
    # exactly once (AC-4). Header/footer paragraph+cell extraction uses
    # _p_text_no_txbx so header-anchored textboxes stay exclusively COM-owned
    # (AC-3; design.md Option C).
    seen_parts: set = set()
    _HF_SLOTS = (
        "header", "footer", "first_page_header",
        "first_page_footer", "even_page_header", "even_page_footer",
    )
    for s_idx, section in enumerate(doc.sections):
        for slot_name in _HF_SLOTS:
            slot = getattr(section, slot_name)
            if slot.is_linked_to_previous:  # optimization only, NOT the guarantee
                continue
            root_el = slot._element  # <w:hdr>/<w:ftr> root (R-1 confirmed)
            if root_el is None or root_el in seen_parts:
                continue  # element-identity dedup = the guarantee
            seen_parts.add(root_el)
            _process_container_content(
                slot, f"HdrFtr[s{s_idx}:{slot_name}]", 1,
                text_extractor=_p_text_no_txbx,
            )

    check_document_size_limits(
        segment_count=len(segs),
        total_text_length=total_text_length,
        max_segments=max_segments,
        max_text_length=max_text_length,
        document_type="Word document",
    )

    return segs


def _insert_docx_translations(
    doc: Any,
    segs: List[Segment],
    tmap: Dict[Tuple, str],
    targets: List[str],
    log: Callable[[str], None] = lambda s: None,
    output_mode: str = "append",
) -> Tuple[int, int]:
    """Insert translations from *tmap* into the document segments.

    tmap key schema (IP-4 / BR-81):
        - table cell segments: (tgt, src_text, col)   — col is 0-based column index
        - non-table segments:  (tgt, src_text, None)  — col=None for paragraphs/textboxes
    """
    ok_cnt = skip_cnt = 0

    def _add_formatted_run(p: Paragraph, text: str, italic: bool, font_size_pt: int) -> None:
        lines = text.split("\n")
        for i, line in enumerate(lines):
            run = p.add_run(line)
            if italic:
                run.italic = True
            if font_size_pt:
                run.font.size = Pt(font_size_pt)
            if i < len(lines) - 1:
                run.add_break()
        tag_run = p.add_run(INSERT_MARKER)
        if italic:
            tag_run.italic = True
        if font_size_pt:
            tag_run.font.size = Pt(font_size_pt)

    for seg in segs:
        # IP-4 (BR-81): tmap key uses (tgt, text, col) for table cells (col=int)
        # and (tgt, text, None) for non-table segments (para/txbx).
        _seg_key = (seg.text, seg.col)  # col is None for non-table segs
        has_any_translation = any((tgt, seg.text, seg.col) in tmap for tgt in targets)
        if not has_any_translation:
            log(f"[SKIP] No translation: {seg.ctx} | {seg.text[:50]}...")
            continue

        translations = []
        for tgt in targets:
            if (tgt, seg.text, seg.col) in tmap:
                translations.append(tmap[(tgt, seg.text, seg.col)])
            else:
                log(f"[WARN] Missing {tgt} translation: {seg.text[:30]}...")
                translations.append(f"[Translation missing|{tgt}] {seg.text[:50]}...")

        if seg.kind == "para":
            if hasattr(seg.ref, "tag") and seg.ref.tag.endswith("}sdt"):
                sdt_element = seg.ref
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                sdt_content = sdt_element.find(qn("w:sdtContent"))
                if sdt_content is not None:
                    if output_mode == "replace":
                        # Overwrite the first paragraph in the SDT content in-place.
                        all_paras = sdt_content.xpath(".//w:p", namespaces=ns)
                        if all_paras:
                            p_obj = Paragraph(all_paras[0], None)
                            runs = p_obj.runs
                            if runs:
                                runs[0].text = translations[0]
                                for r in runs[1:]:
                                    r.text = ""
                            else:
                                p_obj.add_run(translations[0])
                        else:
                            new_p_element = OxmlElement("w:p")
                            sdt_content.append(new_p_element)
                            Paragraph(new_p_element, None).add_run(translations[0])
                        ok_cnt += 1
                        continue
                    existing_paras = sdt_content.xpath(".//w:p", namespaces=ns)
                    existing_texts = []
                    for ep in existing_paras:
                        p_obj = Paragraph(ep, None)
                        if _is_our_insert_block(p_obj):
                            existing_texts.append(_p_text_no_txbx(p_obj))
                    if len(existing_texts) >= len(translations):
                        if all(normalize_text(e) == normalize_text(t) for e, t in zip(existing_texts[:len(translations)], translations)):
                            skip_cnt += 1
                            log(f"[SKIP] SDT translation exists: {seg.text[:30]}...")
                            continue
                    for t in translations:
                        if not any(normalize_text(t) == normalize_text(e) for e in existing_texts):
                            new_p_element = OxmlElement("w:p")
                            sdt_content.append(new_p_element)
                            new_p = Paragraph(new_p_element, None)
                            _add_formatted_run(new_p, t, italic=True, font_size_pt=INSERT_FONT_SIZE_PT)
                    ok_cnt += 1
                    continue
            p: Paragraph = seg.ref
            if isinstance(p._parent, _Cell):
                cell = p._parent
                try:
                    if output_mode == "replace":
                        # Overwrite the paragraph runs in-place.
                        runs = p.runs
                        if runs:
                            runs[0].text = translations[0]
                            for r in runs[1:]:
                                r.text = ""
                        else:
                            p.add_run(translations[0])
                        ok_cnt += 1
                        continue
                    cell_paragraphs = list(cell.paragraphs)
                    p_index = -1
                    for idx, cell_p in enumerate(cell_paragraphs):
                        if cell_p._element == p._element:
                            p_index = idx
                            break
                    if p_index == -1:
                        log("[WARN] Paragraph not found in cell; fallback insertion")
                        for block in translations:
                            new_p = cell.add_paragraph()
                            _add_formatted_run(new_p, block, italic=True, font_size_pt=INSERT_FONT_SIZE_PT)
                        ok_cnt += 1
                        continue
                    existing_texts = []
                    check_limit = min(p_index + 1 + len(translations), len(cell_paragraphs))
                    for idx in range(p_index + 1, check_limit):
                        if _is_our_insert_block(cell_paragraphs[idx]):
                            existing_texts.append(_p_text_no_txbx(cell_paragraphs[idx]))
                    if len(existing_texts) >= len(translations):
                        if all(normalize_text(e) == normalize_text(t) for e, t in zip(existing_texts[:len(translations)], translations)):
                            skip_cnt += 1
                            log(f"[SKIP] Cell already has translations: {seg.text[:30]}...")
                            continue
                    to_add = []
                    for t in translations:
                        if not any(normalize_text(t) == normalize_text(e) for e in existing_texts):
                            to_add.append(t)
                    if not to_add:
                        skip_cnt += 1
                        log(f"[SKIP] Cell translations already present: {seg.text[:30]}...")
                        continue
                    insert_after = p
                    for block in to_add:
                        try:
                            new_p_element = OxmlElement("w:p")
                            insert_after._element.addnext(new_p_element)
                            new_p = Paragraph(new_p_element, cell)
                            _add_formatted_run(new_p, block, italic=True, font_size_pt=INSERT_FONT_SIZE_PT)
                            insert_after = new_p
                        except Exception as exc:
                            log(f"[ERROR] Cell insertion failed: {exc}")
                            try:
                                new_p = cell.add_paragraph()
                                _add_formatted_run(new_p, block, italic=True, font_size_pt=INSERT_FONT_SIZE_PT)
                            except Exception as exc2:
                                log(f"[ERROR] Fallback insertion failed: {exc2}")
                                continue
                    ok_cnt += 1
                except Exception as exc:
                    log(f"[ERROR] Cell processing failed: {exc}")
                    continue
            else:
                try:
                    if output_mode == "replace":
                        # Overwrite run text in-place; only the first translation is used
                        # (multi-target is clamped to append by the orchestrator, BR-67).
                        replacement = translations[0]
                        runs = p.runs
                        if runs:
                            runs[0].text = replacement
                            for r in runs[1:]:
                                r.text = ""
                        else:
                            p.add_run(replacement)
                        ok_cnt += 1
                    elif output_mode == "bilingual":
                        # Emit a two-column, one-row table: col-A = source, col-B = translation.
                        # ADR-0007: source paragraph is replaced by the table; run-level
                        # formatting is not preserved (cell text is a fresh run).
                        source_text = p.text
                        translation = translations[0]
                        # Skip empty paragraphs — pass them through unchanged.
                        if not source_text.strip():
                            ok_cnt += 1
                        else:
                            tbl_obj = doc.add_table(rows=1, cols=2)
                            # Relocate the table from the body end to the paragraph's position.
                            parent = p._element.getparent()
                            p_idx = list(parent).index(p._element)
                            parent.insert(p_idx, tbl_obj._tbl)
                            parent.remove(p._element)
                            tbl_obj.cell(0, 0).text = source_text
                            tbl_obj.cell(0, 1).text = translation
                            ok_cnt += 1
                    else:
                        existing_texts = _scan_our_tail_texts(p, limit=max(len(translations), 4))
                        if existing_texts and len(existing_texts) >= len(translations):
                            if all(normalize_text(e) == normalize_text(t) for e, t in zip(existing_texts[:len(translations)], translations)):
                                skip_cnt += 1
                                log(f"[SKIP] Paragraph already has translations: {seg.text[:30]}...")
                                continue
                        to_add = []
                        for t in translations:
                            if not any(normalize_text(t) == normalize_text(e) for e in existing_texts):
                                to_add.append(t)
                        if not to_add:
                            skip_cnt += 1
                            log(f"[SKIP] Paragraph translations already present: {seg.text[:30]}...")
                            continue
                        last = _find_last_inserted_after(p, limit=max(len(translations), 4))
                        anchor = last if last else p
                        for block in to_add:
                            try:
                                anchor = _append_after(anchor, block, italic=True, font_size_pt=INSERT_FONT_SIZE_PT)
                            except Exception as exc:
                                log(f"[ERROR] Paragraph insertion failed: {exc}")
                                try:
                                    new_p = p._parent.add_paragraph(block)
                                    if new_p.runs:
                                        new_p.runs[0].italic = True
                                except Exception as exc2:
                                    log(f"[ERROR] Fallback insertion failed: {exc2}")
                                    continue
                        ok_cnt += 1
                except Exception as exc:
                    log(f"[ERROR] Paragraph processing failed: {exc}")
                    continue
        elif seg.kind == "txbx":
            tx = seg.ref
            if output_mode == "replace":
                # Overwrite the text in the first paragraph of the text box in-place.
                paras = tx.xpath("./*[local-name()='p']")
                if paras:
                    t_nodes = paras[0].xpath(".//*[local-name()='t']")
                    if t_nodes:
                        t_nodes[0].text = translations[0]
                        for t_node in t_nodes[1:]:
                            t_node.text = ""
                    else:
                        r_elem = OxmlElement("w:r")
                        t_elem = OxmlElement("w:t")
                        t_elem.text = translations[0]
                        r_elem.append(t_elem)
                        paras[0].append(r_elem)
                ok_cnt += 1
            else:
                if _txbx_tail_equals(tx, translations):
                    skip_cnt += 1
                    continue
                paras = tx.xpath("./*[local-name()='p']")
                tail_texts = []
                scan = paras[-max(len(translations), 4):] if len(paras) else []
                for q in scan:
                    has_zero = any(((t.text or "").find(INSERT_MARKER) >= 0) for t in q.xpath(".//*[local-name()='t']"))
                    if has_zero:
                        qtxt = "".join([(node.text or "") for node in q.xpath(".//*[local-name()='t' or local-name()='br']")]).strip()
                        tail_texts.append(qtxt)
                to_add = []
                for t in translations:
                    if not any(normalize_text(t) == normalize_text(e) for e in tail_texts):
                        to_add.append(t)
                if not to_add:
                    skip_cnt += 1
                    continue
                for block in to_add:
                    _txbx_append_paragraph(tx, block, italic=True, font_size_pt=INSERT_FONT_SIZE_PT)
                ok_cnt += 1
        elif seg.kind == "cell":
            # IP-4: whole-cell translation from the serializer path.
            # seg.ref is the _Cell object; translations are keyed by (tgt, text, col).
            cell = seg.ref
            try:
                # Check if translation already inserted (has INSERT_MARKER paragraph)
                if any(_is_our_insert_block(p) for p in cell.paragraphs):
                    skip_cnt += 1
                    log(f"[SKIP] Cell already has translations: {seg.text[:30]}...")
                    continue
                if output_mode == "replace":
                    # Replace first paragraph text in-place with the first translation
                    replacement = translations[0]
                    paras = list(cell.paragraphs)
                    if paras and paras[0].runs:
                        paras[0].runs[0].text = replacement
                        for run in paras[0].runs[1:]:
                            run.text = ""
                    elif paras:
                        paras[0].text = replacement
                    else:
                        new_p = cell.add_paragraph()
                        new_p.add_run(replacement)
                else:
                    # Append each translation as a new italic paragraph in the cell
                    for block in translations:
                        new_p = cell.add_paragraph()
                        _add_formatted_run(new_p, block, italic=True, font_size_pt=INSERT_FONT_SIZE_PT)
                ok_cnt += 1
            except Exception as exc:
                log(f"[ERROR] Cell translation insert failed: {exc}")
                continue
    log(f"[DOCX] Inserted: {ok_cnt} segments, skipped: {skip_cnt}")
    return ok_cnt, skip_cnt


def _translate_docx_via_doc2doc(
    texts: List[str],
    target: str,
    src_lang: Optional[str],
    client,
    stop_flag=None,
    log: Callable[[str], None] = lambda s: None,
    max_batch_chars: int = DEFAULT_MAX_BATCH_CHARS,
    terms=None,
    in_path: str = "",
    status_callback: Optional[Callable[[Optional[str]], None]] = None,
) -> Tuple[Dict[Tuple[str, str], str], int, int, bool]:
    """Use semantic chunking (translate_document) for long single-target DOCX docs."""
    from app.backend.models.translatable_document import (
        TranslatableDocument,
        TranslatableElement,
        ElementType,
        PageInfo,
        DocumentMetadata,
    )
    from app.backend.services.translation_service import translate_document

    elements = [
        TranslatableElement(
            element_id=f"seg-{i}",
            content=t,
            element_type=ElementType.TEXT,
            page_num=1,
        )
        for i, t in enumerate(texts)
    ]
    doc_ir = TranslatableDocument(
        source_path=in_path,
        source_type="docx",
        elements=elements,
        pages=[PageInfo(page_num=1, width=595.0, height=842.0)],
        metadata=DocumentMetadata(),
    )

    try:
        translate_document(
            doc_ir, [target], src_lang, client,
            stop_flag=stop_flag, log=log,
            terms=terms, max_batch_chars=max_batch_chars,
        )
    except RuntimeError as exc:
        log(f"[DOCX] Doc2Doc chunking failed, falling back to batch: {exc}")
        return translate_texts(
            texts, [target], src_lang, client,
            max_batch_chars=max_batch_chars, stop_flag=stop_flag, log=log, terms=terms,
            status_callback=status_callback,
        )

    tmap: Dict[Tuple[str, str], str] = {}
    fail_cnt = 0
    for elem in doc_ir.elements:
        key = (target, elem.content)
        val = elem.translated_content
        if val and not val.startswith("[Translation failed"):
            tmap[key] = val
        else:
            tmap[key] = val or f"[Translation failed|{elem.content[:20]}]"
            fail_cnt += 1
    return tmap, len(texts) - fail_cnt, fail_cnt, False


def _recover_truncated_cell(
    cell_text: str,
    tgt: str,
    src_lang: Optional[str],
    client: OllamaClient,
    max_batch_chars: int,
    stop_flag: Optional[threading.Event],
    log: Callable[[str], None],
) -> str:
    """Recover ONE flagged (suspiciously-short) table cell (truncation-length-
    guard, BR-117, ADR-0020 decision 3).

    Mirrors the BR-82 split-and-retranslate pattern (~L1101-1132) for a
    single cell: split on "\n", re-translate each unique non-empty line via
    `translate_texts`, reassemble. Bounded to ONE attempt and NON-RE-ENTRANT
    — this helper never calls `is_suspiciously_short` on its own output, so
    it cannot loop (ADR-0020 reversal-guarded invariant 2).

    Missing lines (not returned by `translate_texts`) fall back to their
    original line text within the reassembly, matching the existing BR-82
    partial-recovery behavior — this is not a whole-source substitution.
    """
    src_for_prompt = src_lang or "auto"
    lines = cell_text.split("\n")
    uniq_lines = list(dict.fromkeys(
        line for line in lines if line.strip() and should_translate(line, src_for_prompt)
    ))
    fallback_tmap: Dict = {}
    if uniq_lines:
        fallback_tmap, _, _, _ = translate_texts(
            uniq_lines, [tgt], src_lang, client,
            max_batch_chars=max_batch_chars,
            stop_flag=stop_flag, log=log,
        )
    return "\n".join(fallback_tmap.get((tgt, line), line) for line in lines)


def translate_docx(
    in_path: str,
    out_path: str,
    targets: List[str],
    src_lang: Optional[str],
    client: OllamaClient,
    include_headers_shapes_via_com: bool,
    stop_flag: Optional[threading.Event] = None,
    log: Callable[[str], None] = lambda s: None,
    max_batch_chars: int = DEFAULT_MAX_BATCH_CHARS,
    pre_translate_hook: Optional[Callable[[List[str]], None]] = None,
    post_translate_hook: Optional[Callable[[List[Tuple[str, str, str]]], None]] = None,
    terms_getter: Optional[Callable[[], list]] = None,
    output_mode: str = "append",
    block_overrides: Optional[Dict[str, str]] = None,
    status_callback: Optional[Callable[[Optional[str]], None]] = None,
) -> bool:
    from shutil import copyfile

    copyfile(in_path, out_path)
    doc = docx.Document(out_path)

    ok, msg = client.health_check()
    log(f"[API] {msg}")
    if not ok:
        raise ApiError("Ollama service unavailable")

    segs = _collect_docx_segments(doc)
    log(f"[DOCX] segments: {len(segs)}")

    # IP-4: separate non-table segments (para/txbx) from table cell segments.
    # Table cells are translated per-table via the shared serializer (one LLM call
    # per table); non-table segments use the existing translate_texts path.
    para_segs = [s for s in segs if s.table_id is None]
    cell_segs = [s for s in segs if s.table_id is not None]

    # Deduplicate non-table segment texts for translation batch
    seen_texts: set[str] = set()
    uniq_texts: list[str] = []  # kept for post_translate_hook indexing
    for s in para_segs:
        if s.text not in seen_texts and should_translate(s.text, (src_lang or "auto")):
            seen_texts.add(s.text)
            uniq_texts.append(s.text)
    if pre_translate_hook:
        pre_translate_hook(uniq_texts)
    _terms = terms_getter() if terms_getter else None

    # p3-llm-judge: block_overrides seam — when provided, use stored re-translated text
    # instead of calling the LLM (D7). Block ids use the same key as post_translate_hook.
    import os as _os
    file_stem = _os.path.splitext(_os.path.basename(in_path))[0]
    stopped = False

    # para_tmap uses 2-element keys (tgt, text) — standard translate_texts output.
    # final_tmap uses 3-element keys (tgt, text, col) for the restore pass (BR-81).
    para_tmap: Dict = {}
    fail_cnt = 0

    if block_overrides is not None:
        # Build para_tmap from overrides map; applies to non-table segments only.
        for idx, src_text in enumerate(uniq_texts):
            block_id = f"docx:{file_stem}:{idx}"
            if block_id in block_overrides:
                for tgt in targets:
                    para_tmap[(tgt, src_text)] = block_overrides[block_id]
            else:
                for tgt in targets:
                    para_tmap[(tgt, src_text)] = src_text
        log(f"[DOCX] block_overrides applied: {len(block_overrides)} overrides, {len(uniq_texts)} blocks")
    else:
        # Long-document semantic chunking path (P2-6): use translate_document for
        # single-target docs over the char threshold so doc_chunker actually runs.
        _LONG_DOC_CHARS = 40_000
        _total_chars = sum(len(t) for t in uniq_texts)
        if len(targets) == 1 and _total_chars > _LONG_DOC_CHARS:
            para_tmap, _, fail_cnt, stopped = _translate_docx_via_doc2doc(
                uniq_texts, targets[0], src_lang, client,
                stop_flag=stop_flag, log=log,
                max_batch_chars=max_batch_chars, terms=_terms,
                in_path=in_path, status_callback=status_callback,
            )
        else:
            para_tmap, _, fail_cnt, stopped = translate_texts(
                uniq_texts,
                targets,
                src_lang,
                client,
                max_batch_chars=max_batch_chars,
                stop_flag=stop_flag,
                log=log,
                terms=_terms,
                status_callback=status_callback,
            )

        if fail_cnt:
            log(f"[DOCX] failed translations: {fail_cnt}")

        if fail_cnt and not stopped:
            from app.backend.utils.translation_verification import verify_and_fill_tmap
            verify_and_fill_tmap(para_tmap, client, src_lang, stop_flag=stop_flag, log=log)

    # Build 3-element final_tmap from para_tmap (re-key col=None) and table cell tmap.
    final_tmap: Dict = {}
    for (tgt, text), tr in para_tmap.items():
        final_tmap[(tgt, text, None)] = tr

    # IP-4: translate each table group via serializer → one translate_once call per table.
    if cell_segs and not stopped:
        from collections import defaultdict
        from app.backend.utils import table_serializer

        # Group cell segments by table_id
        table_groups: Dict = defaultdict(list)
        for s in cell_segs:
            table_groups[s.table_id].append(s)

        # Duck-typed proxy for table_serializer.serialize() (needs .row/.col/.content/.is_numeric)
        from dataclasses import dataclass as _dc

        @_dc
        class _CellProxy:
            row: int
            col: int
            content: str
            is_numeric: bool = False

        for table_id, t_segs in table_groups.items():
            if stop_flag and stop_flag.is_set():
                stopped = True
                break
            # Determine grid dimensions from segments
            num_rows = max(s.row for s in t_segs) + 1
            num_cols = max(s.col for s in t_segs) + 1
            # Build cells compatible with table_serializer

            # Include ALL grid positions (including empty cells as positional placeholders)
            cells_by_pos: Dict = {(s.row, s.col): s.text for s in t_segs}

            for tgt in targets:
                if stop_flag and stop_flag.is_set():
                    stopped = True
                    break
                src_for_prompt = src_lang or "auto"
                translated_by_pos: Optional[Dict[Tuple[int, int], str]] = None

                if config.JSON_STRUCTURED_TRANSLATION_ENABLED:
                    # IP-5: coordinate JSON envelope (BR-79/BR-80/BR-82) — content-cells only.
                    content_cells = [
                        _CellProxy(row=r, col=c, content=text)
                        for (r, c), text in cells_by_pos.items()
                        if text.strip() and not is_numeric_cell(text)
                    ]
                    if not content_cells:
                        translated_by_pos = {}
                    else:
                        sent_cells = {(c.row, c.col): c.content for c in content_cells}
                        payload = json_translation.build_table_payload(content_cells, src_for_prompt, tgt)
                        try:
                            ok, response = client.translate_json(payload, system_context=None)
                            if ok:
                                translated_by_pos, reason = table_serializer.parse_json(response, sent_cells)
                                if translated_by_pos is None:
                                    logger.warning(
                                        "[DOCX] Table group %s: parse_json() rejected reply for target=%s "
                                        "(%s); falling back to per-cell batch",
                                        table_id, tgt, reason,
                                    )
                                    log(f"[DOCX] Table group {table_id}: JSON table fallback ({reason})")
                            else:
                                logger.warning(
                                    "[DOCX] Table group %s translate_json failed (target=%s): %s; "
                                    "falling back to per-cell batch",
                                    table_id, tgt, response,
                                )
                                log(f"[DOCX] Table group {table_id}: JSON table fallback (translate_json failed)")
                        except Exception as exc:
                            logger.warning(
                                "[DOCX] Table group %s whole-table JSON call raised (target=%s): %s; "
                                "falling back to per-cell batch",
                                table_id, tgt, exc,
                            )
                            log(f"[DOCX] Table group {table_id}: JSON table fallback (exception: {exc})")
                else:
                    # Flag-OFF: retained legacy pipe-grid block, unchanged (Resolution A).
                    proxy_cells = [
                        _CellProxy(row=r, col=c, content=cells_by_pos.get((r, c), ""))
                        for r in range(num_rows) for c in range(num_cols)
                    ]
                    serialized = table_serializer.serialize(proxy_cells)
                    prompt = client._build_table_translate_prompt(serialized, src_for_prompt, tgt)
                    # grid stays None on any error → fallback always runs (BR-82)
                    grid = None
                    try:
                        ok, response = client.translate_once(prompt, tgt, src_lang)
                        if ok:
                            grid = table_serializer.parse(response, num_rows, num_cols)
                            if grid is None:
                                logger.warning(
                                    "[DOCX] Table group %s: parse() returned None (expected %d×%d); "
                                    "falling back to per-cell batch for target=%s",
                                    table_id, num_rows, num_cols, tgt,
                                )
                    except Exception as exc:
                        logger.warning(
                            "[DOCX] Table group %s translate_once failed (target=%s): %s; "
                            "falling back to per-cell batch",
                            table_id, tgt, exc,
                        )
                    if grid is not None:
                        translated_by_pos = {
                            (r, c): grid[r][c]
                            for r in range(num_rows) for c in range(num_cols)
                        }

                if translated_by_pos is not None:
                    # truncation-length-guard (BR-117): a well-formed, shape-valid
                    # reply can still be suspiciously short (recorded 4827->370
                    # bug). Flag per accepted cell, recover ONCE (non-re-entrant,
                    # never re-checked), keep the longest of {accepted, recovered}.
                    # Recovered values are cached by source text so a cell text
                    # repeated across columns is recovered once and logs once.
                    _recovered_cells: Dict[str, str] = {}
                    for s in t_segs:
                        r, c = s.row, s.col
                        if s.text.strip() and (r, c) in translated_by_pos:
                            accepted = translated_by_pos[(r, c)]
                            if is_suspiciously_short(s.text, accepted, tgt):
                                if s.text not in _recovered_cells:
                                    recovered = _recover_truncated_cell(
                                        s.text, tgt, src_lang, client,
                                        max_batch_chars, stop_flag, log,
                                    )
                                    kept = accepted if len(accepted) >= len(recovered) else recovered
                                    logger.warning(
                                        "[DOCX] Table group %s: truncation-guard flagged cell "
                                        "(target=%s) accepted_len=%d recovered_len=%d kept_len=%d",
                                        table_id, tgt, len(accepted), len(recovered), len(kept),
                                    )
                                    _recovered_cells[s.text] = kept
                                final_tmap[(tgt, s.text, c)] = _recovered_cells[s.text]
                            else:
                                final_tmap[(tgt, s.text, c)] = accepted
                else:
                    if stop_flag and stop_flag.is_set():
                        # Table groups processed after the stop signal fires never reach
                        # the translate/fallback logic at all: their cells get ZERO
                        # final_tmap entries (not even a "[Translation failed]" placeholder),
                        # surfacing downstream as an invisible "[SKIP] No translation" with
                        # no error trail. Log this explicitly so it isn't mistaken for a
                        # should_translate() filtering decision.
                        logger.warning(
                            "[DOCX] Table group %s: stop_flag set before fallback ran; "
                            "%d cell(s) will get NO translation for target=%s",
                            table_id, sum(1 for s in t_segs if s.text.strip()), tgt,
                        )
                    # Fallback: translate each unique cell text individually (BR-82)
                    uniq_cell_texts = list(dict.fromkeys(
                        s.text for s in t_segs if s.text.strip()
                        and should_translate(s.text, src_for_prompt)
                    ))
                    excluded_by_should_translate = [
                        s.text for s in t_segs
                        if s.text.strip() and not should_translate(s.text, src_for_prompt)
                    ]
                    if excluded_by_should_translate:
                        logger.info(
                            "[DOCX] Table group %s: %d cell(s) excluded by should_translate() "
                            "(target=%s), e.g. %r",
                            table_id, len(excluded_by_should_translate), tgt,
                            excluded_by_should_translate[0][:30],
                        )
                    if uniq_cell_texts:
                        # A cell can hold MULTIPLE original paragraphs joined with "\n"
                        # (_collect_docx_segments' cell_text_parts join) — most commonly
                        # a merged "layout" cell holding an entire document section.
                        # Sending the whole blob as ONE translate_once() call risks
                        # silent truncation: confirmed live against panjit's gpt-oss:120b,
                        # a 4827-char cell returned only 370 chars with ok=True (no error,
                        # since the response wasn't EMPTY, just cut short) — over 90% of
                        # the content vanished with no trace. Splitting on "\n" and
                        # translating at the SAME per-paragraph granularity as body text
                        # (reusing translate_texts' batching/context/critique) keeps every
                        # individual LLM call bounded, then cells are reassembled by
                        # rejoining their translated lines.
                        cell_to_lines: Dict[str, List[str]] = {
                            cell_text: cell_text.split("\n") for cell_text in uniq_cell_texts
                        }
                        uniq_lines = list(dict.fromkeys(
                            line for lines in cell_to_lines.values() for line in lines
                            if line.strip() and should_translate(line, src_for_prompt)
                        ))
                        fallback_tmap: Dict = {}
                        if uniq_lines:
                            fallback_tmap, _, _, _ = translate_texts(
                                uniq_lines, [tgt], src_lang, client,
                                max_batch_chars=max_batch_chars,
                                stop_flag=stop_flag, log=log,
                            )
                        for cell_text, lines in cell_to_lines.items():
                            translated_lines = [
                                fallback_tmap.get((tgt, line), line) for line in lines
                            ]
                            reassembled = "\n".join(translated_lines)
                            # Apply to ALL cells with this text across their actual columns
                            for s in t_segs:
                                if s.text == cell_text:
                                    final_tmap[(tgt, cell_text, s.col)] = reassembled
                        missing_lines = set(uniq_lines) - {line for (_, line) in fallback_tmap}
                        if missing_lines:
                            logger.warning(
                                "[DOCX] Table group %s: %d cell line(s) sent to fallback "
                                "translate_texts() but missing from its result (target=%s), "
                                "e.g. %r — those lines fall back to their original untranslated "
                                "text within the reassembled cell",
                                table_id, len(missing_lines), tgt, next(iter(missing_lines))[:30],
                            )

    if final_tmap:
        # R1: doc2doc long-doc path always uses append; output_mode="replace" is a follow-up.
        effective_mode = "append" if (len(targets) == 1 and sum(len(t) for t in uniq_texts) > 40_000) else output_mode
        _insert_docx_translations(doc, segs, final_tmap, targets, log=log, output_mode=effective_mode)

    if post_translate_hook is not None:
        tuples: List[Tuple[str, str, str]] = []
        for idx, src_text in enumerate(uniq_texts):
            for tgt in targets:
                if (tgt, src_text) in para_tmap:  # use para_tmap (2-element) for hook
                    tuples.append((f"docx:{file_stem}:{idx}", src_text, para_tmap[(tgt, src_text)]))
        if tuples:
            post_translate_hook(tuples)

    doc.save(out_path)
    if stopped:
        log(f"[DOCX] partial output: {os.path.basename(out_path)}")
    else:
        log(f"[DOCX] output: {os.path.basename(out_path)}")

    if not stopped and include_headers_shapes_via_com and is_win32com_available():
        postprocess_docx_shapes_with_word(out_path, targets, src_lang, client, include_headers=True, log=log)

    return stopped
