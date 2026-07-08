"""PDF translation processor.

Supports two modes:
1. PyMuPDF mode (default): Extracts text with bbox, enables layout preservation
2. PyPDF2 fallback: Simple text extraction for compatibility
"""

from __future__ import annotations

import logging
import os
import threading
from pathlib import Path
from typing import TYPE_CHECKING, Any, Callable, Dict, List, Optional

import docx
from PyPDF2 import PdfReader

from app.backend.clients.ollama_client import OllamaClient
from app.backend.config import (
    MAX_SEGMENTS,
    MAX_TEXT_LENGTH,
    PDF_DRAW_MASK,
    PDF_HEADER_FOOTER_MARGIN_PT,
    PDF_PARSER_ENGINE,
    PDF_SKIP_HEADER_FOOTER,
    PDF_TABLE_ROW_GROWTH_ENABLED,
)
from app.backend.processors.com_helpers import is_win32com_available, word_convert
from app.backend.processors.docx_processor import translate_docx
from app.backend.utils.exceptions import check_document_size_limits
from app.backend.utils.translation_helpers import translate_blocks_batch

if TYPE_CHECKING:
    from app.backend.models.translatable_document import TranslatableDocument
    from app.backend.models.translatable_document import TranslatableDocument

logger = logging.getLogger(__name__)

# pdf-renderer-fallback-warn: exact warning strings (em-dash —, NOT ASCII hyphen)
FITZ_FALLBACK_WARNING = (
    "PDF rendering quality reduced: fell back to basic renderer — "
    "images and formatting may be lost"
)
DOCX_ROUTING_WARNING = (
    "Layout preservation skipped: PDF was converted to bilingual DOCX mode — "
    "use output_format=pdf for layout-faithful output"
)
# BR-104 (pdf-render-truncation-disclosure, AC-11): one aggregated entry per
# file naming the affected page(s) when any element's translated text was
# truncated by the BR-36 fit cascade (render_truncated=True). Disclosure-only —
# never fails the job, never alters output.
TEXT_TRUNCATION_WARNING_TEMPLATE = (
    "'{doc_id}' page(s) {pages}: translated text did not fully fit its layout "
    "box and was truncated — content may be incomplete in the rendered output"
)

def _apply_formula_passthrough(elements: list) -> None:
    """Set FORMULA elements as non-translatable with content as pass-through (AC-7, IP-R4).

    This is path-independent: runs for both the ONNX detector path and the
    heuristic path.  The detector path already sets should_translate=False for
    FORMULA elements; this ensures the heuristic path does the same.

    Mutates elements in-place.

    Args:
        elements: List of TranslatableElement objects; FORMULA-typed ones are updated.
    """
    from app.backend.models.translatable_document import ElementType
    for elem in elements:
        if elem.element_type == ElementType.FORMULA:
            elem.should_translate = False
            if elem.translated_content is None:
                elem.translated_content = elem.content


def _group_table_elements(elements: list) -> Dict[str, list]:
    """Group translatable elements by the table they belong to.

    Elements are grouped by ``metadata["table_id"]`` (assigned by the parser's
    find_tables pass) and must carry a resolved (table_row, table_col) grid
    position.  Groups with fewer than 2 elements are dropped — a single cell
    gains no context from whole-table serialization.

    Args:
        elements: Translatable elements (any subset of a parsed document).

    Returns:
        Dict mapping table_id -> list of elements in that table.
    """
    groups: Dict[str, list] = {}
    for e in elements:
        tid = e.metadata.get("table_id")
        if (
            tid is not None
            and e.metadata.get("table_row") is not None
            and e.metadata.get("table_col") is not None
        ):
            groups.setdefault(tid, []).append(e)
    return {tid: elems for tid, elems in groups.items() if len(elems) >= 2}


def _translate_pdf_tables_with_context(
    table_groups: Dict[str, list],
    tgt: str,
    src_lang: Optional[str],
    client: OllamaClient,
    stop_flag: Optional[threading.Event] = None,
    log: Callable[[str], None] = lambda s: None,
) -> Dict[str, str]:
    """Translate each detected PDF table with full-table context (one LLM call per table).

    Serializes the table grid via ``table_serializer`` and prompts with
    ``client._build_table_translate_prompt`` — the same whole-table wire format
    used by the DOCX/PPTX/XLSX table paths — so every cell is translated with
    its row/column context instead of in isolation.

    Only cells backed by exactly ONE text element are mapped back (a cell with
    multiple elements cannot split one cell translation across bboxes); all
    other cells, and every cell of a table whose whole-table call fails, fall
    back to the regular flatten batch path.

    Args:
        table_groups: Output of _group_table_elements.
        tgt: Target language.
        src_lang: Source language or None (auto).
        client: LLM client.
        stop_flag: Optional cancellation event.
        log: Progress log callback.

    Returns:
        Dict mapping source text (stripped) -> translated text for every cell
        that was successfully translated with table context.
    """
    from dataclasses import dataclass

    from app.backend.utils import table_serializer

    @dataclass
    class _CellProxy:
        row: int
        col: int
        content: str
        is_numeric: bool = False

    out: Dict[str, str] = {}
    for tid, elems in table_groups.items():
        if stop_flag and stop_flag.is_set():
            break

        by_cell: Dict[tuple, list] = {}
        for e in elems:
            key = (e.metadata["table_row"], e.metadata["table_col"])
            by_cell.setdefault(key, []).append(e)

        num_rows = max(r for r, _ in by_cell) + 1
        num_cols = max(c for _, c in by_cell) + 1
        proxy_cells = [
            _CellProxy(
                row=r,
                col=c,
                content=" ".join(
                    el.content.strip() for el in by_cell.get((r, c), [])
                ).strip(),
            )
            for r in range(num_rows)
            for c in range(num_cols)
        ]

        serialized = table_serializer.serialize(proxy_cells)
        prompt = client._build_table_translate_prompt(serialized, src_lang or "auto", tgt)
        grid = None
        try:
            ok, response = client.translate_once(prompt, tgt, src_lang)
            if ok:
                grid = table_serializer.parse(response, num_rows, num_cols)
                if grid is None:
                    logger.warning(
                        "[PDF] table %s: parse() returned None for target=%s "
                        "(expected %d×%d); cells fall back to flatten path. "
                        "Response excerpt: %s",
                        tid, tgt, num_rows, num_cols,
                        response[:120] if response else "",
                    )
        except Exception as exc:
            logger.warning(
                "[PDF] table %s whole-table call failed (target=%s): %s; "
                "cells fall back to flatten path",
                tid, tgt, exc,
            )

        if grid is None:
            continue

        mapped = 0
        for (r, c), cell_elems in by_cell.items():
            if len(cell_elems) != 1:
                continue  # cannot split one cell translation across element bboxes
            src_text = cell_elems[0].content.strip()
            translated = grid[r][c].strip()
            if src_text and translated:
                out[src_text] = translated
                mapped += 1
        log(f"[PDF] table {tid}: {mapped}/{len(by_cell)} cells translated with table context ({tgt})")

    return out


# Lazy import for PyMuPDF parser
_pymupdf_parser = None


def _get_pymupdf_parser():
    """Lazy-load PyMuPDF parser to avoid import errors if not installed."""
    global _pymupdf_parser
    if _pymupdf_parser is None:
        try:
            from app.backend.parsers.pdf_parser import PyMuPDFParser

            _pymupdf_parser = PyMuPDFParser(
                skip_header_footer=PDF_SKIP_HEADER_FOOTER,
                header_footer_margin_pt=PDF_HEADER_FOOTER_MARGIN_PT,
            )
        except ImportError:
            logger.warning("PyMuPDF not installed, using PyPDF2 fallback")
            _pymupdf_parser = False  # Mark as unavailable
    return _pymupdf_parser if _pymupdf_parser else None


def _save_layout_viz(doc: "TranslatableDocument", in_path: str, out_path: str) -> None:
    """Persist layout_viz.json + page thumbnails for the layout viewer overlay.

    Shared by every PDF processing path (DOCX output and PDF-to-PDF output) so
    ``GET /api/jobs/{id}`` can report ``layout_viz_available`` consistently
    regardless of chosen output_format/layout_mode. Non-critical: all failures
    are swallowed so the render/translation pipeline is never blocked by this.

    Multiple PDFs in the same job merge into one layout_viz.json under "files".
    """
    if not doc.layout_viz:
        return
    import json
    import fitz as _fitz  # noqa: PLC0415
    job_dir = Path(out_path).parent.parent
    viz_path = job_dir / "layout_viz.json"
    try:
        existing_files: dict = {}
        if viz_path.exists():
            try:
                existing_files = json.loads(viz_path.read_text(encoding="utf-8")).get("files", {})
            except Exception:
                existing_files = {}
        file_name = Path(in_path).name
        existing_files[file_name] = {
            "file_name": file_name,
            "total_pages": len(doc.layout_viz),
            "pages": doc.layout_viz,
        }
        viz_path.write_text(
            json.dumps({"files": existing_files}, ensure_ascii=False),
            encoding="utf-8",
        )
    except Exception:
        pass  # viz is non-critical; don't block translation

    # Render page thumbnails for layout viewer image overlay.
    # Store per-file under layout_pages/<stem>/page_N.jpg so multi-file
    # jobs don't collide.  Rendering is non-critical; errors are swallowed.
    pages_dir = job_dir / "layout_pages" / Path(in_path).stem
    pages_dir.mkdir(parents=True, exist_ok=True)
    try:
        render_doc = _fitz.open(in_path)
        mat = _fitz.Matrix(1.2, 1.2)  # ~86 DPI — small files, fast render
        for pg in render_doc:
            pix = pg.get_pixmap(matrix=mat)
            jpg_bytes = pix.tobytes(output="jpeg", jpg_quality=60)
            (pages_dir / f"page_{pg.number + 1}.jpg").write_bytes(jpg_bytes)
        render_doc.close()
    except Exception:
        pass  # thumbnails are non-critical


def translate_pdf(
    in_path: str,
    out_path: str,
    targets: List[str],
    src_lang: Optional[str],
    client: OllamaClient,
    stop_flag: Optional[threading.Event] = None,
    log: Callable[[str], None] = lambda s: None,
    use_pymupdf: Optional[bool] = None,
    skip_header_footer: Optional[bool] = None,
    output_format: str = "docx",
    layout_mode: str = "inline",
    draw_mask: Optional[bool] = None,
    pre_translate_hook: Optional[Callable[[List[str]], None]] = None,
    post_translate_hook: Optional[Callable[[List[Tuple[str, str, str]]], None]] = None,
    block_overrides: Optional[Dict[str, str]] = None,
    warnings_callback: Optional[Callable[[str], None]] = None,
    status_callback: Optional[Callable[[Optional[str], Optional[Any]], None]] = None,
) -> bool:
    """Translate a PDF file.

    Args:
        in_path: Input PDF file path.
        out_path: Output file path (DOCX or PDF depending on output_format).
        targets: List of target languages.
        src_lang: Source language (or None for auto-detect).
        cache: Translation cache instance.
        client: Ollama client for translation.
        stop_flag: Optional threading event to signal stop.
        log: Logging callback function.
        use_pymupdf: Force PyMuPDF (True) or PyPDF2 (False). None uses config.
        skip_header_footer: Override header/footer skipping. None uses config.
        output_format: Output format - "docx" (default) or "pdf" (requires Phase 3).
        layout_mode: Layout mode - "inline" (default), "overlay", or "side_by_side".
        draw_mask: Draw white mask over original text in overlay mode. None uses config.
        status_callback: Optional progress hook, called as
            ``status_callback(message, segment=None)`` (pdf-stage-detail-snapshot,
            BR-105 parity) with a ``CurrentSegmentSnapshot(stage="translate", ...)``
            each time a segment finishes translating on the flatten batch path.

    Returns:
        True if stopped early, False if completed.

    Raises:
        ValueError: If output_format="pdf" with layout_mode="inline".
    """
    # Validate output_format and layout_mode combination
    if output_format == "pdf" and layout_mode == "inline":
        raise ValueError(
            "PDF output format is not supported with inline layout mode. "
            "Use layout_mode='overlay' or 'side_by_side' for PDF output, "
            "or use output_format='docx' for inline mode."
        )

    # PDF output with overlay/side_by_side - use PDFGenerator
    if output_format == "pdf" and layout_mode in ("overlay", "side_by_side"):
        return _translate_pdf_to_pdf(
            in_path,
            out_path,
            targets,
            src_lang,
            client,
            stop_flag,
            log,
            skip_header_footer,
            layout_mode,
            draw_mask,
            pre_translate_hook=pre_translate_hook,
            post_translate_hook=post_translate_hook,
            block_overrides=block_overrides,
            warnings_callback=warnings_callback,
            status_callback=status_callback,
        )

    # Try Windows COM conversion first (highest quality)
    temp_docx = str(Path(out_path).with_suffix("")) + "__from_pdf.docx"
    if is_win32com_available():
        try:
            word_convert(in_path, temp_docx, 16)
            from app.backend.processors.docx_processor import translate_docx as _translate_docx
            stopped = _translate_docx(
                temp_docx,
                out_path,
                targets,
                src_lang,
                client,
                include_headers_shapes_via_com=False,
                stop_flag=stop_flag,
                log=log,
                pre_translate_hook=pre_translate_hook,
                post_translate_hook=post_translate_hook,
                block_overrides=block_overrides,
            )
            try:
                os.remove(temp_docx)
            except OSError:
                pass
            return stopped
        except (OSError, RuntimeError) as exc:
            log(f"[PDF] Word import failed, fallback to text extract: {exc}")

    # Notify caller that layout-faithful output is not available (DOCX route only).
    if warnings_callback:
        warnings_callback(DOCX_ROUTING_WARNING)

    # Determine which parser to use
    should_use_pymupdf = use_pymupdf if use_pymupdf is not None else (PDF_PARSER_ENGINE == "pymupdf")
    parser = _get_pymupdf_parser() if should_use_pymupdf else None

    if parser:
        return _translate_pdf_with_pymupdf(
            in_path,
            out_path,
            targets,
            src_lang,
            client,
            stop_flag,
            log,
            skip_header_footer,
            pre_translate_hook=pre_translate_hook,
            post_translate_hook=post_translate_hook,
            block_overrides=block_overrides,
            status_callback=status_callback,
        )
    else:
        return _translate_pdf_with_pypdf2(
            in_path,
            out_path,
            targets,
            src_lang,
            client,
            stop_flag,
            log,
            pre_translate_hook=pre_translate_hook,
            post_translate_hook=post_translate_hook,
            status_callback=status_callback,
            block_overrides=block_overrides,
        )


def _translate_pdf_with_pymupdf(
    in_path: str,
    out_path: str,
    targets: List[str],
    src_lang: Optional[str],
    client: OllamaClient,
    stop_flag: Optional[threading.Event],
    log: Callable[[str], None],
    skip_header_footer: Optional[bool],
    pre_translate_hook: Optional[Callable[[List[str]], None]] = None,
    post_translate_hook: Optional[Callable[[List[Tuple[str, str, str]]], None]] = None,
    block_overrides: Optional[Dict[str, str]] = None,
    status_callback: Optional[Callable[[Optional[str], Optional[Any]], None]] = None,
) -> bool:
    """Translate PDF using PyMuPDF parser with bbox support.

    This method provides:
    - Correct reading order
    - Header/footer detection and optional skipping
    - Table region awareness
    - Better text block handling
    """
    from app.backend.models.translatable_document import ElementType
    from app.backend.parsers.pdf_parser import PyMuPDFParser

    # Create parser with appropriate settings
    should_skip_hf = skip_header_footer if skip_header_footer is not None else PDF_SKIP_HEADER_FOOTER
    parser = PyMuPDFParser(
        skip_header_footer=should_skip_hf,
        header_footer_margin_pt=PDF_HEADER_FOOTER_MARGIN_PT,
    )

    try:
        # Parse PDF
        log(f"[PDF] Parsing with PyMuPDF: {os.path.basename(in_path)}")
        doc = parser.parse(in_path)

        # Save layout viz data if available (non-critical).
        _save_layout_viz(doc, in_path, out_path)

        if not doc.metadata.has_text_layer:
            log("[PDF] Warning: PDF appears to be scanned (low text content)")

        # Get elements in reading order
        elements = doc.get_elements_in_reading_order()

        # FORMULA pass-through (IP-R4): path-independent; heuristic path does not
        # set should_translate=False for FORMULA by default, so apply it here.
        _apply_formula_passthrough(elements)

        translatable = [e for e in elements if e.should_translate and e.content.strip()]

        log(f"[PDF] Found {len(translatable)} translatable blocks across {doc.metadata.page_count} pages")

        # Calculate total text length for size limit check
        total_text_length = sum(len(e.content.strip()) for e in translatable)

        # Check document size limits
        check_document_size_limits(
            segment_count=len(translatable),
            total_text_length=total_text_length,
            max_segments=MAX_SEGMENTS,
            max_text_length=MAX_TEXT_LENGTH,
            document_type="PDF document",
        )

        # IP-7 (p3-table-structure, BR-70): Partition structured-table elements from
        # the flatten batch.  Table elements carrying a recognized TableStructure are
        # translated via the cell-batch seam; all other elements go through the
        # existing flatten translate_blocks_batch path.
        _table_rec_enabled = os.environ.get(
            "TABLE_RECOGNITION_ENABLED", "false"
        ).lower() in ("1", "true", "yes")

        structured_table_elems = []
        flatten_translatable = []
        for e in translatable:
            if (
                _table_rec_enabled
                and e.element_type.value == "table"
                and e.metadata.get("table_structure") is not None
            ):
                structured_table_elems.append(e)
            else:
                flatten_translatable.append(e)

        # Collect unique texts for batch translation (flatten path only — BR-70).
        # dict.fromkeys preserves reading order so the sliding-context prefix in
        # translate_merged_paragraphs sees the REAL neighboring text (set() gave
        # arbitrary neighbors), and block_override IDs are stable across runs.
        unique_texts = list(dict.fromkeys(e.content.strip() for e in flatten_translatable if e.content.strip()))
        log(f"[PDF] Translating {len(unique_texts)} unique texts ({len(structured_table_elems)} structured tables via cell-batch)")
        if pre_translate_hook:
            pre_translate_hook(unique_texts)

        # Batch translate for each target language
        stopped = False
        translations_by_target = {}

        # p3-llm-judge: block_overrides seam — when provided, build translations_by_target
        # from the stored per-block map instead of calling the LLM (D7).
        _file_stem = os.path.splitext(os.path.basename(in_path))[0]
        if block_overrides is not None:
            for tgt in targets:
                translations_by_target[tgt] = {}
                for idx, src_text in enumerate(unique_texts):
                    block_id = f"pdf:{_file_stem}:{idx}"
                    if block_id in block_overrides:
                        translations_by_target[tgt][src_text] = block_overrides[block_id]
                    else:
                        translations_by_target[tgt][src_text] = src_text
            log(f"[PDF] block_overrides applied: {len(block_overrides)} overrides, {len(unique_texts)} unique texts")
        else:
            # Cell-batch seam for structured tables (runs before flatten batch)
            if structured_table_elems:
                try:
                    from app.backend.services.translation_service import translate_table_cells
                    for elem in structured_table_elems:
                        if stop_flag and stop_flag.is_set():
                            break
                        translate_table_cells(
                            element=elem,
                            targets=targets,
                            src_lang=src_lang,
                            client=client,
                            stop_flag=stop_flag,
                            log=log,
                        )
                except Exception as exc:
                    logger.warning(
                        "[PDF] cell-batch seam raised %s; structured tables will fall back to flatten path.",
                        exc,
                    )
                    # Fall back: add them to the flatten path
                    for e in structured_table_elems:
                        if e.content.strip() and e.content.strip() not in set(unique_texts):
                            unique_texts.append(e.content.strip())
                    structured_table_elems = []

            # Whole-table context groups (find_tables-backed, no ONNX needed):
            # cells in these groups are translated one-table-per-call so each
            # cell sees its row/column context.
            table_groups = _group_table_elements(flatten_translatable)
            if table_groups:
                log(f"[PDF] {len(table_groups)} table(s) will be translated with whole-table context")

            for tgt in targets:
                if stop_flag and stop_flag.is_set():
                    log(f"[STOP] PDF stopped before translating to {tgt}")
                    stopped = True
                    break

                table_tmap = (
                    _translate_pdf_tables_with_context(
                        table_groups, tgt, src_lang, client, stop_flag=stop_flag, log=log
                    )
                    if table_groups
                    else {}
                )
                # Texts already translated with table context skip the flatten batch;
                # unmapped table cells (multi-element cells, failed tables) stay in it.
                flatten_texts = [t for t in unique_texts if t not in table_tmap]

                # pdf-stage-detail-snapshot: emit CurrentSegmentSnapshot(stage="translate")
                # per segment on the flatten batch path, mirroring translation_service's
                # status_callback wiring for Office formats (BR-105 parity).
                _on_segment_done = None
                if status_callback is not None:
                    def _on_segment_done(src_text: str, translated: str, _tgt: str = tgt) -> None:
                        from app.backend.services.job_manager import CurrentSegmentSnapshot
                        status_callback(
                            f"翻譯中…（{_tgt}）",
                            CurrentSegmentSnapshot(stage="translate", source=src_text, draft=translated),
                        )

                log(f"[PDF] Batch translating to {tgt}...")
                results = translate_blocks_batch(
                    flatten_texts, tgt, src_lang, client, log=log,
                    on_segment_done=_on_segment_done,
                )
                translations_by_target[tgt] = {
                    text: (translated if ok else f"[Translation failed|{tgt}] {text}")
                    for text, (ok, translated) in zip(flatten_texts, results)
                }
                translations_by_target[tgt].update(table_tmap)
                if not (stop_flag and stop_flag.is_set()):
                    from app.backend.utils.translation_verification import verify_and_fill_dict
                    verify_and_fill_dict(translations_by_target[tgt], tgt, client, src_lang, stop_flag=stop_flag, log=log)

        # Create output document
        output_doc = docx.Document()
        current_page = 0

        for i, element in enumerate(flatten_translatable):
            if stop_flag and stop_flag.is_set():
                log(f"[STOP] PDF stopped at element {i}/{len(flatten_translatable)}")
                stopped = True
                break

            # Add page header when page changes
            if element.page_num != current_page:
                current_page = element.page_num
                output_doc.add_heading(f"-- Page {current_page} --", level=1)

            # Add original text with element type indicator
            type_indicator = ""
            if element.element_type == ElementType.HEADER:
                type_indicator = "[Header] "
            elif element.element_type == ElementType.FOOTER:
                type_indicator = "[Footer] "
            elif element.element_type == ElementType.TABLE_CELL:
                type_indicator = "[Table] "

            original_text = element.content.strip()
            output_doc.add_paragraph(f"{type_indicator}{original_text}")

            # Add translations for each target language
            for tgt in targets:
                if tgt in translations_by_target and original_text in translations_by_target[tgt]:
                    translated = translations_by_target[tgt][original_text]
                else:
                    translated = f"[No translation|{tgt}] {original_text}"
                output_doc.add_paragraph(translated)

            if stopped:
                break

        Path(out_path).parent.mkdir(parents=True, exist_ok=True)
        output_doc.save(out_path)

        if post_translate_hook is not None:
            import os as _os
            file_stem = _os.path.splitext(_os.path.basename(in_path))[0]
            tuples: List[Tuple[str, str, str]] = []
            for idx, src_text in enumerate(unique_texts):
                for tgt in targets:
                    if tgt in translations_by_target and src_text in translations_by_target[tgt]:
                        tuples.append((f"pdf:{file_stem}:{idx}", src_text, translations_by_target[tgt][src_text]))
            if tuples:
                post_translate_hook(tuples)

        if stopped:
            log(f"[PDF] Partial output: {os.path.basename(out_path)}")
        else:
            log(f"[PDF] Output: {os.path.basename(out_path)}")

        return stopped

    except Exception as exc:
        log(f"[PDF] PyMuPDF parsing failed, falling back to PyPDF2: {exc}")
        return _translate_pdf_with_pypdf2(
            in_path, out_path, targets, src_lang, client, stop_flag, log,
            post_translate_hook=post_translate_hook,
            block_overrides=block_overrides,
        )


def _translate_pdf_with_pypdf2(
    in_path: str,
    out_path: str,
    targets: List[str],
    src_lang: Optional[str],
    client: OllamaClient,
    stop_flag: Optional[threading.Event],
    log: Callable[[str], None],
    pre_translate_hook: Optional[Callable[[List[str]], None]] = None,
    post_translate_hook: Optional[Callable[[List[Tuple[str, str, str]]], None]] = None,
    block_overrides: Optional[Dict[str, str]] = None,
    status_callback: Optional[Callable[[Optional[str], Optional[Any]], None]] = None,
) -> bool:
    """Translate PDF using PyPDF2 (fallback method).

    This is the original implementation kept for compatibility.
    Uses batch translation for better context preservation.
    """
    doc = docx.Document()
    stopped = False
    translations_by_target: dict = {}  # initialized before try to prevent UnboundLocalError
    unique_texts: list = []

    try:
        reader = PdfReader(in_path)
        total_pages = len(reader.pages)

        # Extract all page texts first
        page_texts = []
        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            page_texts.append((i, text.strip()))

        # Calculate total text length for size limit check
        total_text_length = sum(len(text) for _, text in page_texts if text)

        # Check document size limits
        check_document_size_limits(
            segment_count=len([t for _, t in page_texts if t]),
            total_text_length=total_text_length,
            max_segments=MAX_SEGMENTS,
            max_text_length=MAX_TEXT_LENGTH,
            document_type="PDF document",
        )

        # Collect unique texts for batch translation (page order preserved for context)
        unique_texts = list(dict.fromkeys(text for _, text in page_texts if text))
        log(f"[PDF] PyPDF2: {total_pages} pages, {len(unique_texts)} unique texts")
        if pre_translate_hook:
            pre_translate_hook(unique_texts)

        # p3-llm-judge: block_overrides seam
        _pypdf2_stem = os.path.splitext(os.path.basename(in_path))[0]
        translations_by_target = {}
        if block_overrides is not None:
            for tgt in targets:
                translations_by_target[tgt] = {}
                for idx, src_text in enumerate(unique_texts):
                    block_id = f"pdf:{_pypdf2_stem}:{idx}"
                    if block_id in block_overrides:
                        translations_by_target[tgt][src_text] = block_overrides[block_id]
                    else:
                        translations_by_target[tgt][src_text] = src_text
            log(f"[PDF/PyPDF2] block_overrides applied: {len(block_overrides)} overrides")
        else:
            # Batch translate for each target language
            for tgt in targets:
                if stop_flag and stop_flag.is_set():
                    log(f"[STOP] PDF stopped before translating to {tgt}")
                    stopped = True
                    break

                # pdf-stage-detail-snapshot: emit CurrentSegmentSnapshot(stage="translate")
                # per segment, mirroring translation_service's status_callback wiring
                # for Office formats (BR-105 parity).
                _on_segment_done = None
                if status_callback is not None:
                    def _on_segment_done(src_text: str, translated: str, _tgt: str = tgt) -> None:
                        from app.backend.services.job_manager import CurrentSegmentSnapshot
                        status_callback(
                            f"翻譯中…（{_tgt}）",
                            CurrentSegmentSnapshot(stage="translate", source=src_text, draft=translated),
                        )

                log(f"[PDF] Batch translating to {tgt}...")
                results = translate_blocks_batch(
                    unique_texts, tgt, src_lang, client, log=log,
                    on_segment_done=_on_segment_done,
                )
                translations_by_target[tgt] = {
                    text: (translated if ok else f"[Translation failed|{tgt}] {text}")
                    for text, (ok, translated) in zip(unique_texts, results)
                }
                if not (stop_flag and stop_flag.is_set()):
                    from app.backend.utils.translation_verification import verify_and_fill_dict
                    verify_and_fill_dict(translations_by_target[tgt], tgt, client, src_lang, stop_flag=stop_flag, log=log)

        # Build output document
        for page_num, text in page_texts:
            if stop_flag and stop_flag.is_set():
                stopped = True
                break

            doc.add_heading(f"-- Page {page_num} --", level=1)
            if text:
                doc.add_paragraph(text)
                for tgt in targets:
                    if tgt in translations_by_target and text in translations_by_target[tgt]:
                        tr = translations_by_target[tgt][text]
                    else:
                        tr = f"[No translation|{tgt}] {text}"
                    doc.add_paragraph(tr)

    except Exception as exc:
        doc.add_paragraph(f"[PDF extract error] {exc}")

    if post_translate_hook is not None and translations_by_target:
        import os as _os
        file_stem = _os.path.splitext(_os.path.basename(in_path))[0]
        tuples: List[Tuple[str, str, str]] = []
        for idx, src_text in enumerate(unique_texts):
            for tgt in targets:
                if tgt in translations_by_target and src_text in translations_by_target[tgt]:
                    tuples.append((f"pdf:{file_stem}:{idx}", src_text, translations_by_target[tgt][src_text]))
        if tuples:
            post_translate_hook(tuples)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    if stopped:
        log(f"[PDF] partial output: {os.path.basename(out_path)}")
    else:
        log(f"[PDF] output: {os.path.basename(out_path)}")

    return stopped


def _translate_pdf_to_pdf(
    in_path: str,
    out_path: str,
    targets: List[str],
    src_lang: Optional[str],
    client: OllamaClient,
    stop_flag: Optional[threading.Event],
    log: Callable[[str], None],
    skip_header_footer: Optional[bool],
    layout_mode: str,
    draw_mask: Optional[bool] = None,
    pre_translate_hook: Optional[Callable[[List[str]], None]] = None,
    post_translate_hook: Optional[Callable[[List[Tuple[str, str, str]]], None]] = None,
    block_overrides: Optional[Dict[str, str]] = None,
    warnings_callback: Optional[Callable[[str], None]] = None,
    status_callback: Optional[Callable[[Optional[str], Optional[Any]], None]] = None,
) -> bool:
    """Translate PDF to PDF with layout preservation.

    Uses PDFGenerator to create overlay or side-by-side output.
    Supports multiple target languages by generating separate PDF files.

    Args:
        in_path: Input PDF path.
        out_path: Output PDF path (base path for multiple languages).
        targets: Target languages.
        src_lang: Source language.
        cache: Translation cache.
        client: Ollama client.
        stop_flag: Stop flag.
        log: Logging callback.
        skip_header_footer: Whether to skip header/footer.
        layout_mode: 'overlay' or 'side_by_side'.
        draw_mask: Whether to draw white mask over original text. None uses config default.

    Returns:
        True if stopped early, False if completed.
    """
    from app.backend.parsers.pdf_parser import PyMuPDFParser
    from app.backend.renderers.base import RenderMode

    # Create parser
    should_skip_hf = skip_header_footer if skip_header_footer is not None else PDF_SKIP_HEADER_FOOTER
    parser = PyMuPDFParser(
        skip_header_footer=should_skip_hf,
        header_footer_margin_pt=PDF_HEADER_FOOTER_MARGIN_PT,
    )

    try:
        # Parse PDF
        log(f"[PDF] Parsing with PyMuPDF: {os.path.basename(in_path)}")
        doc = parser.parse(in_path)

        # Save layout viz data if available (non-critical).
        _save_layout_viz(doc, in_path, out_path)

        if not doc.metadata.has_text_layer:
            log("[PDF] Warning: PDF appears to be scanned (low text content)")

        # Get translatable elements
        elements = doc.get_elements_in_reading_order()
        translatable = [e for e in elements if e.should_translate and e.content.strip()]

        log(f"[PDF] Found {len(translatable)} translatable blocks across {doc.metadata.page_count} pages")

        # Calculate total text length for size limit check
        total_text_length = sum(len(e.content.strip()) for e in translatable)

        # Check document size limits
        check_document_size_limits(
            segment_count=len(translatable),
            total_text_length=total_text_length,
            max_segments=MAX_SEGMENTS,
            max_text_length=MAX_TEXT_LENGTH,
            document_type="PDF document",
        )

        # Collect unique texts for translation (reading order preserved so the
        # sliding-context prefix in translate_merged_paragraphs sees real neighbors)
        # Note: For PDF overlay mode, we translate blocks individually to preserve layout
        # The translation quality is improved by using paragraph-level granularity in translate_blocks_batch
        unique_texts = list(dict.fromkeys(e.content.strip() for e in translatable))
        log(f"[PDF] Translating {len(unique_texts)} unique text blocks")
        if pre_translate_hook:
            pre_translate_hook(unique_texts)

        # Determine render mode
        mode_enum = RenderMode.OVERLAY if layout_mode == "overlay" else RenderMode.SIDE_BY_SIDE

        # Use draw_mask parameter or fall back to config default
        should_draw_mask = draw_mask if draw_mask is not None else PDF_DRAW_MASK

        stopped = False
        output_files = []

        # p3-llm-judge: block_overrides seam
        _pdf2pdf_stem = os.path.splitext(os.path.basename(in_path))[0]

        # Whole-table context groups (find_tables-backed): cells in these groups
        # are translated one-table-per-call so each cell sees its row/column context.
        table_groups = _group_table_elements(translatable)
        if table_groups:
            log(f"[PDF] {len(table_groups)} table(s) will be translated with whole-table context")

        # Generate PDF for each target language
        if len(targets) > 1:
            log(f"[PDF] Generating {len(targets)} PDF files for languages: {', '.join(targets)}")

        for tgt_idx, tgt in enumerate(targets):
            if stop_flag and stop_flag.is_set():
                log(f"[STOP] PDF stopped before translating to {tgt}")
                stopped = True
                break

            # Generate language-specific output path
            if len(targets) > 1:
                out_stem = Path(out_path).stem
                out_dir = Path(out_path).parent
                lang_suffix = f"_{tgt.replace(' ', '_').replace('-', '_')}"
                lang_out_path = str(out_dir / f"{out_stem}{lang_suffix}.pdf")
            else:
                lang_out_path = out_path

            # p3-llm-judge: use override map if provided
            if block_overrides is not None:
                translations = {}
                for idx, src_text in enumerate(unique_texts):
                    block_id = f"pdf:{_pdf2pdf_stem}:{idx}"
                    if block_id in block_overrides:
                        translations[src_text] = block_overrides[block_id]
                    else:
                        translations[src_text] = src_text
                log(f"[PDF/pdf2pdf] block_overrides applied: {len(block_overrides)} overrides")
                missing_count = 0
            else:
                # Translate texts for this language
                log(f"[PDF] [{tgt_idx + 1}/{len(targets)}] Translating to {tgt}...")

                # Whole-table context translation first; mapped cells skip the
                # flatten batch, unmapped ones stay in it as fallback.
                table_tmap = (
                    _translate_pdf_tables_with_context(
                        table_groups, tgt, src_lang, client, stop_flag=stop_flag, log=log
                    )
                    if table_groups
                    else {}
                )
                flatten_texts = [t for t in unique_texts if t not in table_tmap]

                # pdf-stage-detail-snapshot: emit CurrentSegmentSnapshot(stage="translate")
                # per segment, mirroring translation_service's status_callback wiring
                # for Office formats (BR-105 parity).
                _on_segment_done = None
                if status_callback is not None:
                    def _on_segment_done(src_text: str, translated: str, _tgt: str = tgt) -> None:
                        from app.backend.services.job_manager import CurrentSegmentSnapshot
                        status_callback(
                            f"翻譯中…（{_tgt}）",
                            CurrentSegmentSnapshot(stage="translate", source=src_text, draft=translated),
                        )

                results = translate_blocks_batch(
                    flatten_texts, tgt, src_lang, client, log=log,
                    on_segment_done=_on_segment_done,
                )

                # Build text -> translation mapping
                translations = {}
                missing_count = 0
                for text, (ok, translated) in zip(flatten_texts, results):
                    if ok:
                        translations[text] = translated
                    else:
                        translations[text] = f"[翻譯失敗] {text[:30]}..."
                        missing_count += 1
                translations.update(table_tmap)

                if missing_count > 0:
                    log(f"[PDF] Warning: {missing_count} texts failed to translate to {tgt}")
                    if not (stop_flag and stop_flag.is_set()):
                        from app.backend.utils.translation_verification import verify_and_fill_dict
                        verify_and_fill_dict(translations, tgt, client, src_lang, stop_flag=stop_flag, log=log)

            if post_translate_hook is not None:
                tuples: List[Tuple[str, str, str]] = []
                for element in translatable:
                    content = element.content.strip()
                    if content in translations:
                        translated_text = translations[content]
                        # Only emit real translations (skip failure-placeholder entries)
                        if not translated_text.startswith("[翻譯失敗]"):
                            tuples.append((element.element_id, content, translated_text))
                if tuples:
                    post_translate_hook(tuples)

            # Generate PDF for this language (fitz primary / ReportLab fallback per BR-34)
            _dispatch_render(
                doc=doc,
                translations=translations,
                output_path=lang_out_path,
                target_lang=tgt,
                mode=mode_enum,
                draw_mask=should_draw_mask,
                doc_id=os.path.basename(in_path),
                log=log,
                warnings_callback=warnings_callback,
            )
            output_files.append(lang_out_path)
            log(f"[PDF] Generated: {Path(lang_out_path).name}")

        if stopped:
            log(f"[PDF] Partial output: {len(output_files)} of {len(targets)} files generated")
        else:
            log(f"[PDF] Completed: {len(output_files)} PDF file(s) generated")

        return stopped

    except Exception as exc:
        log(f"[PDF] PDF-to-PDF generation failed: {exc}")
        # Fallback to DOCX output
        log("[PDF] Falling back to DOCX output")
        docx_out = str(Path(out_path).with_suffix(".docx"))
        return _translate_pdf_with_pymupdf(
            in_path, docx_out, targets, src_lang, client, stop_flag, log, skip_header_footer,
            post_translate_hook=post_translate_hook,
        )


# ---------------------------------------------------------------------------
# BR-34 dispatch helpers: fitz primary / ReportLab fallback
# ---------------------------------------------------------------------------

def _run_fitz_render(
    doc: "TranslatableDocument",
    translations: dict,
    output_path: str,
    target_lang: str,
    mode,
    draw_mask: bool,
    log,
) -> None:
    """Invoke the fitz primary renderer.

    Separated as a named function so tests can patch it independently
    of the ReportLab fallback path.
    """
    from app.backend.renderers.fitz_renderer import PDFGenerator

    generator = PDFGenerator(target_lang=target_lang, draw_mask=draw_mask, log=log)
    generator.generate(doc, translations, output_path, mode)


def _run_reportlab_render(
    doc: "TranslatableDocument",
    translations: dict,
    output_path: str,
    target_lang: str,
    mode,
    draw_mask: bool,
    log,
) -> None:
    """Invoke the ReportLab fallback renderer.

    Separated as a named function so tests can patch it independently
    of the fitz primary path.
    """
    from app.backend.renderers.coordinate_renderer import CoordinateRenderer

    renderer = CoordinateRenderer(
        target_lang=target_lang,
        draw_background=draw_mask,
        log=log,
    )
    renderer.render(doc, output_path, translations, mode)


def _dispatch_render(
    doc: "TranslatableDocument",
    translations: dict,
    output_path: str,
    target_lang: str,
    mode,
    draw_mask: bool,
    doc_id: str,
    log=None,
    warnings_callback: Optional[Callable[[str], None]] = None,
) -> None:
    """Dispatch to fitz primary renderer; fall back to ReportLab on unhandled exception.

    Implements BR-34, Table K:
    - fitz primary is always attempted first.
    - On any unhandled exception (import failure, redaction failure, corrupt source):
        * WARNING is logged with exception type + document id.
        * ReportLab fallback is invoked.
    - If ReportLab also raises, the exception propagates (job → failed).
    - Never catches BaseException or KeyboardInterrupt (Decision B).

    AC-10/BR-103: before dispatch, a bounded local table-row-growth pre-pass
    mutates ``doc`` in place (grows over-full table rows within their own
    table/page budget), gated by PDF_TABLE_ROW_GROWTH_ENABLED.
    AC-11/BR-104: after the render call returns (either branch), ``doc`` is
    swept for residual render_truncated markers and exactly one aggregated
    warning is emitted per file via ``warnings_callback`` when any are found.
    """
    if log is None:
        log = lambda s: None  # noqa: E731

    if doc is not None and PDF_TABLE_ROW_GROWTH_ENABLED:
        try:
            from app.backend.renderers.text_region_renderer import grow_table_rows
            grow_table_rows(doc)
        except Exception as exc:
            logger.warning(f"[PDF] row-growth pre-pass failed for '{doc_id}': {exc}")

    try:
        _run_fitz_render(
            doc=doc,
            translations=translations,
            output_path=output_path,
            target_lang=target_lang,
            mode=mode,
            draw_mask=draw_mask,
            log=log,
        )
    except Exception as exc:
        logger.warning(
            f"[PDF] fitz render failed ({type(exc).__name__}) for '{doc_id}'; "
            f"falling back to ReportLab. Exception: {exc}"
        )
        if warnings_callback:
            warnings_callback(FITZ_FALLBACK_WARNING)
        # ReportLab fallback — if this also raises, propagate to job manager.
        _run_reportlab_render(
            doc=doc,
            translations=translations,
            output_path=output_path,
            target_lang=target_lang,
            mode=mode,
            draw_mask=draw_mask,
            log=log,
        )

    _emit_truncation_disclosure_warning(doc, doc_id, warnings_callback)


def _emit_truncation_disclosure_warning(
    doc: "TranslatableDocument",
    doc_id: str,
    warnings_callback: Optional[Callable[[str], None]],
) -> None:
    """AC-11/BR-104: emit exactly one aggregated job.warnings entry per file
    naming the affected page(s) when any element has render_truncated=True.

    Disclosure-only: never raises, never alters output, no-op when doc is
    None/empty or no element was truncated.
    """
    if doc is None or not warnings_callback:
        return
    try:
        truncated_pages = sorted({
            elem.page_num for elem in doc.elements if getattr(elem, "render_truncated", False)
        })
    except Exception:
        return
    if not truncated_pages:
        return
    pages_str = ", ".join(str(p) for p in truncated_pages)
    warnings_callback(TEXT_TRUNCATION_WARNING_TEMPLATE.format(doc_id=doc_id, pages=pages_str))


def parse_pdf_to_document(
    file_path: str,
    skip_header_footer: Optional[bool] = None,
) -> "TranslatableDocument":
    """Parse a PDF file to TranslatableDocument (for advanced use).

    This function exposes the parsed document structure for custom processing.

    Args:
        file_path: Path to the PDF file.
        skip_header_footer: Whether to mark header/footer as non-translatable.

    Returns:
        TranslatableDocument with extracted elements and bbox info.

    Raises:
        ImportError: If PyMuPDF is not installed.
        FileNotFoundError: If file does not exist.
    """
    from app.backend.parsers.pdf_parser import PyMuPDFParser

    should_skip = skip_header_footer if skip_header_footer is not None else PDF_SKIP_HEADER_FOOTER
    parser = PyMuPDFParser(
        skip_header_footer=should_skip,
        header_footer_margin_pt=PDF_HEADER_FOOTER_MARGIN_PT,
    )
    return parser.parse(file_path)
