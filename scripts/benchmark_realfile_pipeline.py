#!/usr/bin/env python3
"""Real-file pipeline benchmark: run actual process_files() end-to-end.

Tests real documents through the full translation pipeline including:
- Context detection (LLM samples doc, describes it)
- Dynamic scenario strategy (auto-detect scenario, apply options/prompt)
- Qwen context flow (append detected context to system prompt)
- Actual file processors (docx, xlsx)
- Batch translation with production prompt builders

Test files:
  - QC-OC151-01 SMD-AU弯脚尺寸量测作业指导卡.xlsx
  - SMD-OI30-D1 SMD Auto Mold作业指导书.doc

Settings (5 model/profile combos × 2 target languages = 10 translation jobs per file):
  1. qwen_general       — Qwen + general profile (pipeline auto-detects scenario)
  2. qwen_techprocess   — Qwen + technical_process profile (forced scenario)
  3. hymt_techprocess   — HY-MT + technical_process profile
  4. hymt_bare          — HY-MT + hymt profile (no system prompt)
  5. tgemma_general     — TranslateGemma + general profile
"""
from __future__ import annotations

import os

os.environ["TRANSLATION_CACHE_ENABLED"] = "0"

import json
import random
import requests
import shutil
import sys
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


def _detect_and_set_ollama_url() -> None:
    """Auto-detect Ollama URL (handles WSL2 gateway) and set env var."""
    if os.environ.get("OLLAMA_BASE_URL"):
        return
    for cand in ("http://localhost:11434", "http://127.0.0.1:11434"):
        try:
            if requests.get(f"{cand}/api/tags", timeout=3).status_code == 200:
                os.environ["OLLAMA_BASE_URL"] = cand
                return
        except Exception:
            pass
    # WSL2: try gateway IP
    try:
        with open("/proc/net/route") as f:
            for line in f:
                parts = line.split()
                if len(parts) >= 3 and parts[1] == "00000000":
                    h = parts[2]
                    ip = f"{int(h[6:8],16)}.{int(h[4:6],16)}.{int(h[2:4],16)}.{int(h[0:2],16)}"
                    cand = f"http://{ip}:11434"
                    if requests.get(f"{cand}/api/tags", timeout=3).status_code == 200:
                        os.environ["OLLAMA_BASE_URL"] = cand
                        return
    except Exception:
        pass


# Detect Ollama URL BEFORE importing backend config (which reads OLLAMA_BASE_URL)
_detect_and_set_ollama_url()

from app.backend.config import TimeoutConfig
from app.backend.processors.orchestrator import process_files
from app.backend.translation_profiles import PROFILES, get_profile

# ╔══════════════════════════════════════════════════════════════════╗
# ║  Constants                                                      ║
# ╚══════════════════════════════════════════════════════════════════╝

SCRIPTS_DIR = Path(__file__).resolve().parent
OUTPUT_BASE = PROJECT_ROOT / "test_output" / "benchmark_realfile_pipeline"

TEST_FILES = [
    SCRIPTS_DIR / "SMD-OI30-D1 SMD Auto Mold作业指导书.doc",
    SCRIPTS_DIR / "QC-OC151-01 SMD-AU弯脚尺寸量测作业指导卡.xlsx",
]

TARGET_LANGUAGES = ["English", "Vietnamese"]

# ╔══════════════════════════════════════════════════════════════════╗
# ║  Settings                                                       ║
# ╚══════════════════════════════════════════════════════════════════╝


@dataclass(frozen=True)
class PipelineSetting:
    name: str
    model_id: str
    model_type: str       # "general" | "translation"
    profile_id: str
    description: str
    num_ctx_override: Optional[int] = None  # primary model ctx (reduce for large models)
    refine_model: Optional[str] = None      # cross-model refiner (None = disabled)
    refiner_num_ctx: Optional[int] = None   # num_ctx override for refiner


SETTINGS = [
    # ── Qwen 9b (candidate upgrade) ──────────────────────────────────
    PipelineSetting(
        name="qwen9b_general",
        model_id="qwen3.5:9b",
        model_type="general",
        profile_id="general",
        description="Qwen 9b + general profile (2048 ctx for VRAM safety)",
        num_ctx_override=2048,
    ),
    PipelineSetting(
        name="qwen9b_techprocess",
        model_id="qwen3.5:9b",
        model_type="general",
        profile_id="technical_process",
        description="Qwen 9b + technical_process profile (2048 ctx for VRAM safety)",
        num_ctx_override=2048,
    ),
    # ── HY-MT + Qwen 4b refiner (current two-phase) ──────────────────
    PipelineSetting(
        name="hymt_qwen4b_refine",
        model_id="demonbyron/HY-MT1.5-7B:Q4_K_M",
        model_type="translation",
        profile_id="technical_process",
        description="HY-MT + Qwen 4b cross-model refiner",
        refine_model="qwen3.5:4b",
    ),
    # ── HY-MT + Qwen 9b refiner (candidate upgrade) ──────────────────
    PipelineSetting(
        name="hymt_qwen9b_refine",
        model_id="demonbyron/HY-MT1.5-7B:Q4_K_M",
        model_type="translation",
        profile_id="technical_process",
        description="HY-MT + Qwen 9b cross-model refiner (2048 ctx for VRAM safety)",
        refine_model="qwen3.5:9b",
        refiner_num_ctx=2048,
    ),
]

# ╔══════════════════════════════════════════════════════════════════╗
# ║  Segment extraction (for comparison report)                     ║
# ╚══════════════════════════════════════════════════════════════════╝


def _extract_docx_segments(path: Path) -> List[str]:
    """Extract translatable text segments from a DOCX file."""
    try:
        import docx as python_docx
        doc = python_docx.Document(str(path))
        segments = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) >= 2:
                segments.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and len(text) >= 2:
                        segments.append(text)
        return segments
    except Exception as e:
        print(f"  Warning: failed to extract DOCX segments: {e}")
        return []


def _extract_xlsx_segments(path: Path) -> List[str]:
    """Extract translatable text segments from an XLSX file."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True, data_only=True)
        segments = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        text = str(cell).strip()
                        if text and len(text) >= 2:
                            segments.append(text)
        wb.close()
        return segments
    except Exception as e:
        print(f"  Warning: failed to extract XLSX segments: {e}")
        return []


def _extract_translated_segments(path: Path) -> List[str]:
    """Extract translated segments from output file."""
    ext = path.suffix.lower()
    if ext == ".docx":
        return _extract_docx_segments(path)
    elif ext == ".xlsx":
        return _extract_xlsx_segments(path)
    return []


# ╔══════════════════════════════════════════════════════════════════╗
# ║  Comparison report                                              ║
# ╚══════════════════════════════════════════════════════════════════╝


def generate_sampling_report(
    all_results: Dict[str, Dict[str, dict]],
    outdir: Path,
    n_samples: int = 10,
    seed: int = 42,
) -> None:
    """Generate a side-by-side sampling comparison report.

    Args:
        all_results: {setting_name: {file_target_key: {segments, source_segments, ...}}}
        outdir: Output directory for the report.
        n_samples: Number of segments to sample per file.
        seed: Random seed for reproducibility.
    """
    lines: List[str] = []
    w = lines.append

    w("# Real-File Pipeline Benchmark — Sampling Comparison")
    w("")
    w(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    w(f"Samples per file/target: {n_samples}")
    w("")

    # Group results by file+target
    file_target_groups: Dict[str, Dict[str, List[str]]] = {}
    source_segments_map: Dict[str, List[str]] = {}

    for setting_name, file_results in all_results.items():
        for ft_key, data in file_results.items():
            if ft_key not in file_target_groups:
                file_target_groups[ft_key] = {}
                source_segments_map[ft_key] = data.get("source_segments", [])
            file_target_groups[ft_key][setting_name] = data.get("translated_segments", [])

    for ft_key in sorted(file_target_groups.keys()):
        w(f"## {ft_key}")
        w("")

        source_segs = source_segments_map.get(ft_key, [])
        setting_segs = file_target_groups[ft_key]

        if not source_segs:
            w("*(No source segments extracted)*")
            w("")
            continue

        # Sample indices
        rng = random.Random(seed)
        n = min(n_samples, len(source_segs))
        indices = sorted(rng.sample(range(len(source_segs)), n))

        for idx in indices:
            src = source_segs[idx]
            w(f"### Segment #{idx + 1}")
            w("")
            w(f"**原文**: {src}")
            w("")

            setting_names = sorted(setting_segs.keys())
            w("| Setting | Translation |")
            w("|---------|------------|")
            for sn in setting_names:
                segs = setting_segs[sn]
                # Try to find matching translated segment
                # The translated file may have different segment count
                # due to inserted translations, so we search for source text match
                translation = _find_translation_for_source(src, segs)
                display = translation[:200] + "..." if len(translation) > 200 else translation
                display = display.replace("|", "\\|").replace("\n", " ")
                w(f"| {sn} | {display} |")
            w("")

    report_path = outdir / "sampling_comparison.md"
    report_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"\nSampling report: {report_path}")


def _find_translation_for_source(source: str, translated_segments: List[str]) -> str:
    """Find the translation corresponding to a source segment.

    In the output docx/xlsx, translations are inserted after the source text.
    We look for the segment immediately after the source match.
    """
    for i, seg in enumerate(translated_segments):
        if seg.strip() == source.strip():
            # Next segment should be the translation
            if i + 1 < len(translated_segments):
                next_seg = translated_segments[i + 1].strip()
                if next_seg and next_seg != source.strip():
                    return next_seg
    # Fallback: return empty
    return "(not found)"


# ╔══════════════════════════════════════════════════════════════════╗
# ║  Main execution                                                 ║
# ╚══════════════════════════════════════════════════════════════════╝


def main() -> None:
    # Validate test files exist
    for f in TEST_FILES:
        if not f.exists():
            sys.exit(f"Test file not found: {f}")

    outdir = OUTPUT_BASE / time.strftime("%Y%m%d_%H%M%S")
    outdir.mkdir(parents=True, exist_ok=True)

    print("=" * 65)
    print("Real-File Pipeline Benchmark")
    print(f"Output: {outdir}")
    print(f"Test files: {len(TEST_FILES)}")
    print(f"Settings: {len(SETTINGS)}")
    print(f"Target languages: {TARGET_LANGUAGES}")
    print(f"Total jobs: {len(SETTINGS) * len(TEST_FILES) * len(TARGET_LANGUAGES)}")
    print("=" * 65)

    # Extract source segments for comparison
    source_segments: Dict[str, List[str]] = {}
    for test_file in TEST_FILES:
        ext = test_file.suffix.lower()
        if ext in (".xlsx", ".xls"):
            source_segments[test_file.name] = _extract_xlsx_segments(test_file)
        elif ext in (".doc", ".docx"):
            # For .doc, we can't extract directly; we'll extract from the converted output
            source_segments[test_file.name] = []
        print(f"  Source segments from {test_file.name}: {len(source_segments[test_file.name])}")

    all_results: Dict[str, Dict[str, dict]] = {}
    summary_rows: List[dict] = []
    t_start = time.time()
    job_idx = 0
    total_jobs = len(SETTINGS) * len(TEST_FILES) * len(TARGET_LANGUAGES)

    # Collect logs for each job
    log_lines: Dict[str, List[str]] = {}

    for setting in SETTINGS:
        setting_dir = outdir / setting.name
        setting_dir.mkdir(parents=True, exist_ok=True)
        all_results[setting.name] = {}

        profile = get_profile(setting.profile_id)

        for tgt_lang in TARGET_LANGUAGES:
            tgt_short = tgt_lang[:2].lower()
            tgt_dir = setting_dir / tgt_short
            tgt_dir.mkdir(parents=True, exist_ok=True)

            for test_file in TEST_FILES:
                job_idx += 1
                job_key = f"{setting.name}_{test_file.stem}_{tgt_short}"
                ft_key = f"{test_file.name} | zh→{tgt_short}"

                print(f"\n{'='*65}")
                print(f"[{job_idx}/{total_jobs}] {setting.name} | {test_file.name} | → {tgt_lang}")
                print(f"  Profile: {setting.profile_id} | Model: {setting.model_id}")

                # Collect logs
                job_logs: List[str] = []

                def log_fn(msg: str, _logs=job_logs) -> None:
                    _logs.append(msg)
                    print(f"    {msg}")

                t0 = time.time()
                try:
                    processed, total, stopped, client = process_files(
                        files=[test_file],
                        output_dir=tgt_dir,
                        targets=[tgt_lang],
                        src_lang=None,  # auto-detect
                        include_headers_shapes_via_com=False,
                        ollama_model=setting.model_id,
                        model_type=setting.model_type,
                        system_prompt=profile.system_prompt,
                        profile_id=setting.profile_id,
                        num_ctx_override=setting.num_ctx_override,
                        timeout_config=TimeoutConfig(),
                        log=log_fn,
                        refine_model=setting.refine_model,
                        refiner_num_ctx=setting.refiner_num_ctx,
                    )
                    elapsed = time.time() - t0

                    # Find output file
                    output_files = list(tgt_dir.glob("*_translated*"))
                    output_file = output_files[0] if output_files else None

                    # Extract translated segments
                    translated_segs = []
                    if output_file:
                        translated_segs = _extract_translated_segments(output_file)

                    # Get source segments (for .doc, try extracting from output docx)
                    src_segs = source_segments.get(test_file.name, [])
                    if not src_segs and output_file:
                        # For .doc files, source segments are interleaved in the output .docx
                        src_segs = _extract_docx_segments(output_file)
                        # Take every other segment (source/translation pairs)
                        if len(src_segs) > 1:
                            src_segs = src_segs[::2]  # Even indices = source

                    all_results[setting.name][ft_key] = {
                        "source_segments": src_segs,
                        "translated_segments": translated_segs,
                        "output_file": str(output_file) if output_file else None,
                    }

                    row = {
                        "setting": setting.name,
                        "model": setting.model_id,
                        "profile": setting.profile_id,
                        "file": test_file.name,
                        "target": tgt_lang,
                        "processed": processed,
                        "total": total,
                        "stopped": stopped,
                        "elapsed_s": round(elapsed, 1),
                        "output_file": str(output_file) if output_file else None,
                        "n_source_segments": len(src_segs),
                        "n_translated_segments": len(translated_segs),
                    }
                    summary_rows.append(row)
                    log_lines[job_key] = job_logs

                    status = "OK" if processed == total and not stopped else "PARTIAL"
                    print(f"  → {status} in {elapsed:.1f}s, "
                          f"output: {output_file.name if output_file else 'NONE'}")

                    # Unload model to free VRAM for next setting
                    if client:
                        client.unload_model()

                except Exception as exc:
                    elapsed = time.time() - t0
                    print(f"  → ERROR in {elapsed:.1f}s: {exc}")
                    summary_rows.append({
                        "setting": setting.name,
                        "model": setting.model_id,
                        "profile": setting.profile_id,
                        "file": test_file.name,
                        "target": tgt_lang,
                        "processed": 0,
                        "total": 1,
                        "stopped": False,
                        "elapsed_s": round(elapsed, 1),
                        "error": str(exc),
                    })
                    log_lines[job_key] = job_logs

    total_time = time.time() - t_start

    # ── Save results ──────────────────────────────────────────────
    summary_path = outdir / "summary.json"
    summary_path.write_text(json.dumps(summary_rows, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"\nSummary: {summary_path}")

    # Save logs
    logs_path = outdir / "logs.json"
    logs_path.write_text(json.dumps(log_lines, indent=2, ensure_ascii=False), encoding="utf-8")

    # Generate sampling comparison report
    generate_sampling_report(all_results, outdir)

    # Generate summary table
    lines: List[str] = []
    w = lines.append
    w("# Real-File Pipeline Benchmark Summary")
    w("")
    w(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    w(f"Total time: {total_time:.0f}s ({total_time/60:.1f} min)")
    w("")
    w("## Settings")
    w("")
    w("| # | Name | Model | Refiner | Profile | Description |")
    w("|---|------|-------|---------|---------|-------------|")
    for i, s in enumerate(SETTINGS, 1):
        refiner = s.refine_model or "—"
        if s.refiner_num_ctx and s.refine_model:
            refiner += f" (ctx={s.refiner_num_ctx})"
        w(f"| {i} | {s.name} | {s.model_id} | {refiner} | {s.profile_id} | {s.description} |")
    w("")
    w("## Results")
    w("")
    w("| Setting | File | Target | Time(s) | Status | Segments |")
    w("|---------|------|--------|---------|--------|----------|")
    for r in summary_rows:
        status = "ERROR" if "error" in r else ("OK" if r.get("processed") == r.get("total") else "PARTIAL")
        segs = r.get("n_translated_segments", "?")
        w(f"| {r['setting']} | {r['file'][:30]} | {r['target']} | {r['elapsed_s']} | {status} | {segs} |")
    w("")

    report_path = outdir / "report.md"
    report_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"Report: {report_path}")

    print(f"\nDone! {len(summary_rows)} jobs, {total_time:.0f}s total ({total_time/60:.1f} min)")


if __name__ == "__main__":
    main()
