import docx

def build(nrows=60, ncols=5):
    d = docx.Document()
    t = d.add_table(rows=nrows, cols=ncols)
    for r in range(nrows):
        for c in range(ncols):
            t.cell(r,c).text = f"r{r}c{c}"
    return t

# A) id() key, NO explicit gc — pure refcounting, i.e. exactly production
t = build()
seen, coll = set(), 0
for row in t.rows:
    for cell in row.cells:
        k = id(cell._tc)
        if k in seen: coll += 1
        seen.add(k)
print("A) id() key, no explicit gc  -> distinct:", len(seen), " false collisions:", coll)

# B) store the ELEMENT itself (set holds a strong ref, keeps proxy alive)
t = build()
seen, coll = set(), 0
for row in t.rows:
    for cell in row.cells:
        e = cell._tc
        if e in seen: coll += 1
        seen.add(e)
print("B) element key               -> distinct:", len(seen), " false collisions:", coll)

# C) merged-cell dedup must STILL work under B
d = docx.Document()
t2 = d.add_table(rows=2, cols=4)
t2.cell(0,0).merge(t2.cell(0,3)).text = "FRAME"
for c in range(4): t2.cell(1,c).text = f"d{c}"
seen, emitted = set(), []
for r_i, row in enumerate(t2.rows):
    for c_i, cell in enumerate(row.cells):
        e = cell._tc
        if e in seen: continue
        seen.add(e)
        emitted.append((r_i, c_i, cell.text))
print("C) emitted:", emitted)
print("C) emitted count:", len(emitted), "(want 5: 1 merged at col0 + 4 data)")
