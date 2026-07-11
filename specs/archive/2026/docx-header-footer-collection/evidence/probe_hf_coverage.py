import sys, re, glob; sys.path.insert(0,"/home/egg/Projects/Translate_Tool")
import docx
from app.backend.processors.docx_processor import _collect_docx_segments
norm=lambda s: re.sub(r"\s+","",s)
for f in sorted(glob.glob("/home/egg/Projects/Translate_Tool/docs/TEST_DOC/*.docx")):
    d=docx.Document(f)
    # ground truth: every header/footer paragraph text (excluding txbxContent, matching native domain)
    hf_texts=[]
    for sec in d.sections:
        for part in [sec.header,sec.first_page_header,sec.even_page_header,
                     sec.footer,sec.first_page_footer,sec.even_page_footer]:
            for p in part.paragraphs:
                t="".join(x.text or "" for x in p._p.xpath(".//*[local-name()='t' and not(ancestor::*[local-name()='txbxContent'])]")).strip()
                if t: hf_texts.append(t)
            for tbl in part.tables:
                for row in tbl.rows:
                    for c in row.cells:
                        t=c.text.strip()
                        if t: hf_texts.append(t)
    # dedup ground truth to unique
    gt=set(norm(t) for t in hf_texts)
    segs=_collect_docx_segments(d)
    nblob=norm("\n".join(s.text for s in segs))
    missing=[t for t in gt if t not in nblob]
    name=f.rsplit("/",1)[1][:32]
    print(f"{name:34s} hf_unique_texts={len(gt):3d}  missing_from_collection={len(missing)}")
    for m in missing[:3]: print("   MISSING:", repr(m[:50]))
