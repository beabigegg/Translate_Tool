import sys; sys.path.insert(0,"/home/egg/Projects/Translate_Tool")
import docx
d=docx.Document()
sec=d.sections[0]
# header with a paragraph AND a table (mirrors real docs)
sec.header.paragraphs[0].text="HDR_PARA_ZH"
htbl=sec.header.add_table(rows=1, cols=2, width=docx.shared.Inches(4))
htbl.cell(0,0).text="HDR_CELL_A"; htbl.cell(0,1).text="HDR_CELL_B"
sec.footer.paragraphs[0].text="FTR_PARA_ZH"

hdr=sec.header
print("header type:", type(hdr).__name__)
print("header._element tag:", hdr._element.tag.split('}')[-1])
print("header has .paragraphs:", hasattr(hdr,"paragraphs"), " .tables:", hasattr(hdr,"tables"))
# does header._element iterate children like body?
kids=[c.tag.split('}')[-1] for c in hdr._element]
print("header._element children:", kids)
# write-back: set a run's text, does it persist on save?
hdr.paragraphs[0].runs[0].text = "HDR_TRANSLATED"
import tempfile, os
tmp=tempfile.mktemp(suffix=".docx"); d.save(tmp)
d2=docx.Document(tmp)
print("write-back persists:", d2.sections[0].header.paragraphs[0].text=="HDR_TRANSLATED")
os.unlink(tmp)
