import sys; sys.path.insert(0,"/home/egg/Projects/Translate_Tool")
import docx, glob
# What does a section expose? headers/footers can be default/first-page/even
for f in sorted(glob.glob("/home/egg/Projects/Translate_Tool/docs/TEST_DOC/*.docx")):
    d=docx.Document(f)
    print("###", f.rsplit("/",1)[1][:34], "sections:", len(d.sections))
    for i,sec in enumerate(d.sections):
        parts=[("header",sec.header),("first_page_header",sec.first_page_header),
               ("even_page_header",sec.even_page_header),
               ("footer",sec.footer),("first_page_footer",sec.first_page_footer),
               ("even_page_footer",sec.even_page_footer)]
        for name,part in parts:
            npar=sum(1 for p in part.paragraphs if p.text.strip())
            ntbl=len(part.tables)
            linked=getattr(part,"is_linked_to_previous",None)
            tbl_txt=0
            for t in part.tables:
                for r in t.rows:
                    for c in r.cells:
                        if c.text.strip(): tbl_txt+=1
            if npar or ntbl:
                print(f"   sec{i} {name:18s} paras={npar} tables={ntbl}(cells_with_text={tbl_txt}) linked_to_prev={linked}")
